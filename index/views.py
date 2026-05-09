from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login as django_login, logout as django_logout
from django.http import JsonResponse, HttpResponse, HttpResponseForbidden
from django.contrib.auth.decorators import login_required
from django.db.models import Q, Count, Sum, Max, Min  # ДОБАВИТЬ Max и Min
from django.utils import timezone
import os
import time  
from .utils import (
    generate_smeta_pdf, 
    generate_zayavka_pdf,  
    generate_budget_smeta_pdf,
    generate_budget_pismo_pdf,
    generate_fbp_smeta_pdf, 
    generate_fbp_pismo_pdf,
    FONT_BOLD,
    FONT_NAME,
    canvas,
    colors,
    ImageReader,
    io,
    os,
    A4
)
from .forms import (
    SimpleLoginForm, ПользовательФорма, ОтделФорма, ОрганизацияФорма,
    OrganizationLoginForm, ЗаявкаФорма, ОжидаемаяСуммаФорма, ОтчетФорма, ОтчетПланировщикаФорма, СметаПроектФорма, СметаЗатратФорма,
    СметаРассмотрениеФорма, СоздатьЗаявкиИзСметыФорма, ГенерацияДокументаФорма, СметаДанныеФорма, ДокументЗаявкиФорма, СметаБюджетаФорма, СтатьяСметыФорма, ВыборГодаДляСметыФорма, ВыборСметыАдминистрацииФорма
)
from .models import Пользователь, Отдел, Организация, UserProfile, Заявка, ВидыВыплат, ОтчетПланировщика, СметаЗатрат, СметаПроект, СгенерированныйДокумент, ДокументСметы, ДокументЗаявки, СметаФБП, СтатьяСметыФБП, ДокументСметыФБП, ОтчетПоЗаявкам
from .decorators import role_required, planning_required
from django.views.decorators.csrf import csrf_exempt
from django.contrib import messages
from .models import Выплаты, ОжидаемаяСуммаВыплат, ОтчетВыплат
from django.db import models
from docx import Document
from docx.shared import Inches
import io
import json
from datetime import datetime
from django.contrib.admin.views.decorators import staff_member_required
from django.conf import settings
import zipfile
from io import BytesIO
from django.core.mail import EmailMessage
from django.conf import settings
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from django.core.files.base import ContentFile
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Для работы без GUI

def отправить_смету_на_почту(смета, pdf_smeta_data, pdf_pismo_data):
    """Отправка сметы на почту"""
    print("=== ОТПРАВКА НА ПОЧТУ ===")
    
    # Кому отправлять
    to_email = 'kburajikina@mail.ru'  # ваш email
    
    # Тема письма
    subject = f'📄 Смета ФБП на {смета.год} год'
    
    # Текст письма
    message = f"""
    Здравствуйте!
    
    Одобрена новая смета ФБП на {смета.год} год.
    
    📊 ИНФОРМАЦИЯ О СМЕТЕ:
    • Год: {смета.год}
    • Общая сумма: {смета.общая_сумма:,.2f} руб.
    • Количество статей: {смета.статьи.count()}
    • Создатель: {смета.создатель.Фамилия} {смета.создатель.Имя}
    
    📎 ВО ВЛОЖЕНИИ:
    1. Смета расходов (PDF)
    2. Письмо-заявка (PDF)
    
    Для просмотра сметы войдите в систему.
    
    Ссылка: http://127.0.0.1:8000/budget/fbp-smeti/
    
    --
    Автоматическое уведомление от системы ФБП Лениногорск
    """
    
    try:
        # Создаем письмо
        email = EmailMessage(
            subject=subject,
            body=message,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[to_email],
        )
        
        # Добавляем PDF сметы
        email.attach(
            f'Смета_ФБП_{смета.год}.pdf', 
            pdf_smeta_data, 
            'application/pdf'
        )
        
        # Добавляем PDF письма
        email.attach(
            f'Письмо_заявка_{смета.год}.pdf', 
            pdf_pismo_data, 
            'application/pdf'
        )
        
        # Отправляем
        email.send()
        
        print("✅ Письмо отправлено успешно")
        return True
        
    except Exception as e:
        print(f"❌ Ошибка отправки письма: {e}")
        import traceback
        traceback.print_exc()
        return False

# ==================== ГЛАВНАЯ СТРАНИЦА ====================
def welcome_page(request):
    """Главная страница выбора типа входа"""
    return render(request, 'index/welcome.html')


def custom_login_view(request):
    """Вход для сотрудников и администраторов"""
    error = ''
    
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '')
        
        print(f"=== ПОПЫТКА ВХОДА ===")
        print(f"Логин: {username}")
        
        # СНАЧАЛА пробуем стандартную аутентификацию Django (для суперюзеров)
        user = authenticate(request, username=username, password=password)
        
        if user is not None and user.is_superuser:
            print(f"✅ Суперюзер Django: {user.username}")
            django_login(request, user)
            return redirect('home')  # Перенаправляем на страницу суперюзера
        
        # ЕСЛИ НЕ суперюзер, пробуем кастомную аутентификацию
        try:
            user_custom = Пользователь.objects.get(Логин=username)
            print(f"Найден кастомный пользователь: {user_custom.Фамилия} {user_custom.Имя}")
            print(f"Отдел: {user_custom.Отдел.Название if user_custom.Отдел else 'Нет отдела'}")
            
            if user_custom.Пароль == password:
                print("✅ Пароль верный!")
                
                # Очищаем сессию
                request.session.flush()
                
                if user_custom.Отдел and user_custom.Отдел.Название == 'Отдел по планированию и анализу':
                    request.session['user_id'] = user_custom.id
                    request.session['user_role'] = 'planning'
                    request.session['user_name'] = f"{user_custom.Фамилия} {user_custom.Имя}"
                    print("🎯 Перенаправляем на planning_dashboard")
                    return redirect('planning_dashboard')
                
                else:
                    # ВСЕ остальные пользователи (включая бюджетный отдел) идут на user_dashboard
                    request.session['user_id'] = user_custom.id
                    request.session['user_role'] = 'user' 
                    request.session['user_name'] = f"{user_custom.Фамилия} {user_custom.Имя}"
                    print("→ Перенаправляем на user_dashboard")
                    return redirect('user_dashboard')
                    
            else:
                error = 'Неверный пароль'
                print("❌ Пароль неверный")
                
        except Пользователь.DoesNotExist:
            error = 'Пользователь не найден'
            print("❌ Пользователь не найден")
    
    return render(request, 'index/login_simple.html', {'error': error})

def logout_view(request):
    """Выход для сотрудников"""
    django_logout(request)
    request.session.flush()
    return redirect('login')

# ==================== АУТЕНТИФИКАЦИЯ ОРГАНИЗАЦИЙ ====================
def organization_login_view(request):
    """Вход для организаций"""
    
    # Если уже авторизованы как организация - сразу на дашборд
    if request.session.get('organization_authenticated'):
        return redirect('organization_dashboard')
    
    error = ''
    
    if request.method == 'POST':
        login_input = request.POST.get('login', '').strip()
        password_input = request.POST.get('password', '')
        
        try:
            org = Организация.objects.get(Логин=login_input)
            
            if org.check_password(password_input):
                # ОЧИЩАЕМ сессию полностью перед установкой новых данных
                request.session.flush()
                
                # Устанавливаем ТОЛЬКО данные организации
                request.session['organization_id'] = org.id
                request.session['organization_authenticated'] = True
                request.session['user_role'] = 'organization'
                request.session['org_name'] = org.ПолноеНаименование
                
                return redirect('organization_dashboard')
            else:
                error = 'Неверный пароль'
                
        except Организация.DoesNotExist:
            error = 'Организация не найдена'
    
    return render(request, 'index/organization_login.html', {'error': error})

def organization_dashboard(request):
    """Дашборд организации"""
    
    # Проверяем аутентификацию организации через сессию
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    org_id = request.session.get('organization_id')
    if not org_id:
        return redirect('organization_login')
    
    try:
        org = Организация.objects.get(id=org_id)
    except Организация.DoesNotExist:
        request.session.flush()
        return redirect('organization_login')
    
    # Получаем заявки организации
    applications = Заявка.objects.filter(организация=org).order_by('-дата_подачи')
    
    # Статистика для дашборда
    total_applications = applications.count()
    new_applications = applications.filter(статус='new').count()
    approved_applications = applications.filter(статус='approved').count()
    
    context = {
        'organization': org,
        'applications': applications[:5],
        'total_applications': total_applications,
        'new_applications': new_applications,
        'approved_applications': approved_applications,
    }
    
    return render(request, 'index/organization_dashboard.html', context)

def organization_logout(request):
    """Выход организации"""
    request.session.flush()
    return redirect('organization_login')


# ==================== ДЕКОРАТОР ДЛЯ КАСТОМНЫХ ПОЛЬЗОВАТЕЛЕЙ ====================
def custom_user_required(view_func):
    def _wrapped_view(request, *args, **kwargs):
        # Проверяем, что пользователь авторизован как кастомный пользователь
        if not request.session.get('user_id'):
            return redirect('login')
        return view_func(request, *args, **kwargs)
    return _wrapped_view

# ==================== ЛИЧНЫЕ КАБИНЕТЫ ====================
@custom_user_required
def user_dashboard(request):
    """Личный кабинет кастомного пользователя"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    # Получаем ВСЕ заявки (базовый запрос)
    applications = Заявка.objects.all().order_by('-дата_подачи')
    
    # Получаем параметры фильтрации из GET
    status_filter = request.GET.get('status', '')
    payment_type_filter = request.GET.get('payment_type', '')
    date_from_filter = request.GET.get('date_from', '')
    date_to_filter = request.GET.get('date_to', '')
    
    # Сохраняем исходный запрос для статистики (без фильтра по статусу и виду выплаты)
    # НО с учетом дат, если они выбраны
    stats_applications = Заявка.objects.all()
    
    # Применяем фильтр по датам ДЛЯ СТАТИСТИКИ
    if date_from_filter:
        try:
            date_from = datetime.strptime(date_from_filter, '%Y-%m-%d').date()
            stats_applications = stats_applications.filter(дата_подачи__date__gte=date_from)
        except ValueError:
            pass
    
    if date_to_filter:
        try:
            date_to = datetime.strptime(date_to_filter, '%Y-%m-%d').date()
            stats_applications = stats_applications.filter(дата_подачи__date__lte=date_to)
        except ValueError:
            pass
    
    # Применяем фильтры ДЛЯ ОТОБРАЖЕНИЯ ЗАЯВОК В ТАБЛИЦЕ
    if status_filter:
        applications = applications.filter(статус=status_filter)
    
    if payment_type_filter:
        applications = applications.filter(вид_выплаты_id=payment_type_filter)
    
    if date_from_filter:
        try:
            date_from = datetime.strptime(date_from_filter, '%Y-%m-%d').date()
            applications = applications.filter(дата_подачи__date__gte=date_from)
        except ValueError:
            pass
    
    if date_to_filter:
        try:
            date_to = datetime.strptime(date_to_filter, '%Y-%m-%d').date()
            applications = applications.filter(дата_подачи__date__lte=date_to)
        except ValueError:
            pass
    
    # Получаем все виды выплат для фильтра
    payment_types = ВидыВыплат.objects.all()
    
    # СТАТИСТИКА ПО ЗАЯВКАМ (учитывая выбранные даты)
    stats = {
        'new_count': stats_applications.filter(статус='new').count(),
        'approved_count': stats_applications.filter(статус='approved').count(),
        'rejected_count': stats_applications.filter(статус='rejected').count(),
        'total_count': stats_applications.count(),
        'total_amount': stats_applications.aggregate(total=Sum('запрашиваемая_сумма'))['total'] or 0,
    }
    
    # Проверяем отдел пользователя
    is_budget_department = user.Отдел and user.Отдел.Название == 'Бюджетный отдел'
    is_planning_department = user.Отдел and user.Отдел.Название == 'Отдел по планированию и анализу'
    
    context = {
        'user': user,
        'applications': applications,
        'payment_types': payment_types,
        'stats': stats,
        'is_budget_department': is_budget_department,
        'is_planning_department': is_planning_department,
    }
    
    return render(request, 'index/user_dashboard.html', context)
    
@staff_member_required
def home(request):
    """Главная страница суперпользователя - только для staff/superuser"""
    users = Пользователь.objects.all()
    departments = Отдел.objects.all()

    фамилия_поиск = request.GET.get('search_familiya', '').strip()
    отдел_поиск = request.GET.get('search_otdel', '')

    if фамилия_поиск:
        users = users.filter(Фамилия__icontains=фамилия_поиск)
    if отдел_поиск:
        users = users.filter(Отдел_id=отдел_поиск)

    context = {
        'users': users,
        'departments': departments,
        'search_familiya': фамилия_поиск,
        'search_otdel': отдел_поиск,
    }
    return render(request, 'index/home.html', context)

@custom_user_required
def budget_page(request):
    """Страница бюджетного отдела"""
    # Проверяем что пользователь из бюджетного отдела
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    return render(request, 'index/budget_page.html')


# ==================== УПРАВЛЕНИЕ ПОЛЬЗОВАТЕЛЯМИ ====================
def user_create(request):
    if request.method == 'POST':
        form = ПользовательФорма(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return JsonResponse({'success': True})
        else:
            return JsonResponse({'success': False, 'errors': form.errors})
    else:
        form = ПользовательФорма()
    return render(request, 'index/user_form.html', {'form': form})

def user_update(request, pk):
    user = get_object_or_404(Пользователь, pk=pk)
    if request.method == 'POST':
        form = ПользовательФорма(request.POST, request.FILES, instance=user)
        if form.is_valid():
            form.save()
            return JsonResponse({'success': True})
        else:
            return JsonResponse({'success': False, 'errors': form.errors})
    else:
        form = ПользовательФорма(instance=user)
    return render(request, 'index/user_form.html', {'form': form, 'user': user})

def user_delete(request, pk):
    user = get_object_or_404(Пользователь, pk=pk)
    if request.method == 'POST':
        user.delete()
        return JsonResponse({'success': True})
    return JsonResponse({'success': False})


# ==================== УПРАВЛЕНИЕ ОТДЕЛАМИ ====================
@staff_member_required
def отделы_list(request):
    """Список отделов - только для суперюзеров"""
    departments = Отдел.objects.all()
    return render(request, 'index/отделы_list.html', {'departments': departments})

@staff_member_required
def отдел_create(request):
    """Создание отдела - только для суперюзеров"""
    if request.method == 'POST':
        form = ОтделФорма(request.POST)
        if form.is_valid():
            form.save()
            return redirect('отделы_list')
    else:
        form = ОтделФорма()
    return render(request, 'index/отдел_form.html', {'form': form})


def отдел_update(request, pk):
    department = get_object_or_404(Отдел, pk=pk)
    if request.method == 'POST':
        form = ОтделФорма(request.POST, instance=department)
        if form.is_valid():
            form.save()
            return redirect('отделы_list')
    else:
        form = ОтделФорма(instance=department)
    return render(request, 'index/отдел_form.html', {'form': form})

def отдел_delete(request, pk):
    department = get_object_or_404(Отдел, pk=pk)
    if request.method == 'POST':
        department.delete()
        return redirect('отделы_list')
    return render(request, 'index/отдел_confirm_delete.html', {'department': department})


# ==================== УПРАВЛЕНИЕ ОРГАНИЗАЦИЯМИ ====================
@staff_member_required
def организации_list(request):
    """Список организаций - только для суперюзеров"""
    organizations = Организация.objects.all()
    return render(request, 'index/организации_list.html', {'organizations': organizations})

@staff_member_required
def организация_create(request):
    """Создание организации - только для суперюзеров"""
    if request.method == 'POST':
        form = ОрганизацияФорма(request.POST)
        if form.is_valid():
            org = form.save(commit=False)
            raw_password = form.cleaned_data.get('Пароль')
            if raw_password:
                org.set_password(raw_password)
            org.save()
            return redirect('организации_list')
    else:
        form = ОрганизацияФорма()
    return render(request, 'index/организация_form.html', {'form': form})


def организация_update(request, pk):
    organization = get_object_or_404(Организация, pk=pk)
    if request.method == 'POST':
        form = ОрганизацияФорма(request.POST, instance=organization)
        if form.is_valid():
            org = form.save(commit=False)
            raw_password = form.cleaned_data.get('Пароль')
            if raw_password:
                org.set_password(raw_password)
            org.save()
            return redirect('организации_list')
    else:
        form = ОрганизацияФорма(instance=organization)
    return render(request, 'index/организация_form.html', {'form': form})

def организация_delete(request, pk):
    organization = get_object_or_404(Организация, pk=pk)
    if request.method == 'POST':
        organization.delete()
        return redirect('организации_list')
    return render(request, 'index/организация_confirm_delete.html', {'organization': organization})


# ==================== ЗАЯВКИ ОРГАНИЗАЦИЙ ====================
def application_create(request):
    """Создание заявки организацией с возможностью прикрепить документы"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')

    org_id = request.session.get('organization_id')
    org = Организация.objects.get(id=org_id)

    if request.method == 'POST':
        form = ЗаявкаФорма(request.POST)
        if form.is_valid():
            application = form.save(commit=False)
            application.организация = org
            application.полное_наименование = org.ПолноеНаименование
            application.инн = org.ИНН
            application.счет_организации = org.СчетОрганизации

            if form.cleaned_data.get('использовать_другой_вид'):
                application.вид_выплаты = None
            else:
                application.другой_вид_выплаты = None

            application.save()
            
            # Обработка прикрепленных документов
            документы = request.FILES.getlist('documents')  # получаем список файлов
            count = 0
            
            for файл in документы:
                ДокументЗаявки.objects.create(
                    заявка=application,
                    тип='other',
                    название=файл.name,
                    файл=файл,
                    описание=f"Загружен при создании заявки"
                )
                count += 1
            
            messages.success(request, f'Заявка успешно создана. Прикреплено документов: {count}')
            return redirect('organization_application_detail', pk=application.pk)
    else:
        form = ЗаявкаФорма()

    context = {
        'form': form,
        'organization': org,
        'current_date': timezone.now().strftime('%d.%m.%Y'),
    }
    return render(request, 'index/application_create.html', context)

def organization_applications(request):
    """Список заявок организации"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')

    org_id = request.session.get('organization_id')
    org = Организация.objects.get(id=org_id)

    applications = Заявка.objects.filter(организация=org).order_by('-дата_подачи')

    context = {
        'organization': org,
        'applications': applications,
    }
    return render(request, 'index/organization_applications.html', context)


# ==================== УПРАВЛЕНИЕ ЗАЯВКАМИ (БЮДЖЕТНЫЙ ОТДЕЛ) ====================
@login_required
def budget_applications(request):
    """Заявки для бюджетного отдела"""
    try:
        profile = request.user.userprofile
        if not profile.Отдел or profile.Отдел.Название != 'Бюджетный отдел':
            return redirect('home')
    except UserProfile.DoesNotExist:
        return redirect('home')

    new_applications = Заявка.objects.filter(статус='new').order_by('-дата_подачи')
    processed_applications = Заявка.objects.filter(~Q(статус='new')).order_by('-дата_рассмотрения')

    context = {
        'new_applications': new_applications,
        'processed_applications': processed_applications,
    }
    return render(request, 'index/budget_applications.html', context)

@custom_user_required
@csrf_exempt
def application_approve(request, pk):
    """Одобрение заявки для кастомных пользователей"""
    application = get_object_or_404(Заявка, pk=pk)
    application.статус = 'approved'
    application.дата_рассмотрения = timezone.now()
    application.save()
    return redirect('user_dashboard')



@custom_user_required
@csrf_exempt
def application_reject(request, pk):
    """Отклонение заявки для кастомных пользователей с комментарием"""
    application = get_object_or_404(Заявка, pk=pk)
    
    if request.method == 'POST':
        # Проверяем, это AJAX запрос с комментарием или обычная форма
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            try:
                data = json.loads(request.body)
                комментарий = data.get('комментарий', '').strip()
                
                if not комментарий:
                    return JsonResponse({'success': False, 'error': 'Укажите причину отклонения'})
                
                application.статус = 'rejected'
                application.комментарий_отклонения = комментарий
                application.дата_рассмотрения = timezone.now()
                application.дата_отклонения = timezone.now()
                application.save()
                
                return JsonResponse({'success': True})
                
            except json.JSONDecodeError:
                return JsonResponse({'success': False, 'error': 'Ошибка обработки данных'})
        else:
            # Старая логика для обратной совместимости
            application.статус = 'rejected'
            application.дата_рассмотрения = timezone.now()
            application.дата_отклонения = timezone.now()
            application.save()
            return redirect('user_dashboard')
    
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

    
@login_required
def application_detail(request, pk):
    """Детальный просмотр заявки сотрудником"""
    try:
        profile = request.user.userprofile
        if not profile.Отдел or profile.Отдел.Название != 'Бюджетный отдел':
            return redirect('home')
    except UserProfile.DoesNotExist:
        return redirect('home')

    application = get_object_or_404(Заявка, pk=pk)
    
    context = {
        'application': application,
    }
    return render(request, 'index/application_detail.html', context)


@custom_user_required
def user_application_detail(request, pk):
    """Детальный просмотр заявки для обычного пользователя"""
    application = get_object_or_404(Заявка, pk=pk)
    
    context = {
        'application': application,
    }
    return render(request, 'index/user_application_detail.html', context)

# ==================== УПРАВЛЕНИЕ ЗАЯВКАМИ ДЛЯ КАСТОМНЫХ ПОЛЬЗОВАТЕЛЕЙ ====================
def user_application_approve(request, pk):
    """Одобрение заявки для кастомных пользователей"""
    # Проверяем, что пользователь авторизован как кастомный пользователь
    if not request.session.get('user_id'):
        return redirect('login')
    
    application = get_object_or_404(Заявка, pk=pk)
    application.статус = 'approved'
    application.дата_рассмотрения = timezone.now()
    application.save()
    return redirect('user_dashboard')

@custom_user_required
@csrf_exempt
def user_application_reject(request, pk):
    """Отклонение заявки для кастомных пользователей с комментарием"""
    # Проверяем, что пользователь авторизован как кастомный пользователь
    if not request.session.get('user_id'):
        return JsonResponse({'success': False, 'error': 'Требуется авторизация'})
    
    application = get_object_or_404(Заявка, pk=pk)
    
    if request.method == 'POST':
        # Проверяем, это AJAX запрос с комментарием или обычная форма
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            try:
                data = json.loads(request.body)
                комментарий = data.get('комментарий', '').strip()
                
                if not комментарий:
                    return JsonResponse({'success': False, 'error': 'Укажите причину отклонения'})
                
                application.статус = 'rejected'
                application.комментарий_отклонения = комментарий
                application.дата_рассмотрения = timezone.now()
                application.дата_отклонения = timezone.now()
                application.save()
                
                return JsonResponse({'success': True})
                
            except json.JSONDecodeError:
                return JsonResponse({'success': False, 'error': 'Ошибка обработки данных'})
        else:
            # Старая логика для обратной совместимости
            application.статус = 'rejected'
            application.дата_рассмотрения = timezone.now()
            application.дата_отклонения = timezone.now()
            application.save()
            return redirect('user_dashboard')
    
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

# УДАЛИТЕ старую функцию user_application_reject которая была без комментариев

# ==================== УПРАВЛЕНИЕ ЗАЯВКАМИ ДЛЯ БЮДЖЕТНОГО ОТДЕЛА ====================
def user_application_approve(request, pk):
    """Одобрение заявки для кастомных пользователей"""
    # Проверяем, что пользователь авторизован как кастомный пользователь
    if not request.session.get('user_id'):
        return redirect('login')
    
    application = get_object_or_404(Заявка, pk=pk)
    application.статус = 'approved'
    application.дата_рассмотрения = timezone.now()
    application.save()
    return redirect('user_dashboard')

@login_required
def budget_application_approve(request, pk):
    """Одобрение заявки для бюджетного отдела"""
    try:
        profile = request.user.userprofile
        if not profile.Отдел or profile.Отдел.Название != 'Бюджетный отдел':
            return redirect('home')
    except UserProfile.DoesNotExist:
        return redirect('home')

    application = get_object_or_404(Заявка, pk=pk)
    application.статус = 'approved'
    application.дата_рассмотрения = timezone.now()
    application.save()
    return redirect('budget_applications')

@login_required
@csrf_exempt
def budget_application_reject(request, pk):
    """Отклонение заявки для бюджетного отдела с комментарием"""
    try:
        profile = request.user.userprofile
        if not profile.Отдел or profile.Отдел.Название != 'Бюджетный отдел':
            return redirect('home')
    except UserProfile.DoesNotExist:
        return redirect('home')

    application = get_object_or_404(Заявка, pk=pk)
    
    if request.method == 'POST':
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            try:
                data = json.loads(request.body)
                комментарий = data.get('комментарий', '').strip()
                
                if not комментарий:
                    return JsonResponse({'success': False, 'error': 'Укажите причину отклонения'})
                
                application.статус = 'rejected'
                application.комментарий_отклонения = комментарий
                application.дата_рассмотрения = timezone.now()
                application.дата_отклонения = timezone.now()
                application.save()
                
                return JsonResponse({'success': True})
                
            except json.JSONDecodeError:
                return JsonResponse({'success': False, 'error': 'Ошибка обработки данных'})
        else:
            application.статус = 'rejected'
            application.дата_рассмотрения = timezone.now()
            application.дата_отклонения = timezone.now()
            application.save()
            return redirect('budget_applications')
    
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

# ДОБАВЬТЕ ЭТУ ФУНКЦИЮ В views.py
def create_payments_from_approved_applications():
    """Автоматически создает выплаты из принятых заявок"""
    approved_applications = Заявка.objects.filter(статус='approved', выплаты__isnull=True)
    
    for application in approved_applications:
        # Проверяем, не существует ли уже выплата для этой заявки
        if not Выплаты.objects.filter(заявка=application).exists():
            Выплаты.objects.create(
                полное_наименование=application.полное_наименование,
                инн=application.инн,
                счет_получателя=application.счет_получателя,
                сумма=application.запрашиваемая_сумма,
                вид=application.вид_выплаты,
                заявка=application,
                назначение_платежа=f"Выплата по заявке #{application.id}"
            )
            
@custom_user_required
def выплаты_list(request):
    """Список выплат с фильтрацией по организациям"""
    # Автоматически создаем выплаты из принятых заявок
    create_payments_from_approved_applications()
    
    выплаты = Выплаты.objects.all().order_by('-дата_создания')
    
    # Получаем все организации для фильтра
    все_организации = Организация.objects.all()
    
    # Получаем параметры фильтрации
    date_filter = request.GET.get('date_filter', '')
    start_date_str = request.GET.get('start_date', '')
    end_date_str = request.GET.get('end_date', '')
    организации_ids = request.GET.getlist('организации')  # Множественный выбор
    
    # Определяем период для расчета
    сегодня = timezone.now().date()
    start_date = None
    end_date = None
    
    # Флаг, показывающий применена ли фильтрация по дате
    date_filter_applied = False
    
    if date_filter:
        date_filter_applied = True
        if date_filter == 'today':
            start_date = сегодня
            end_date = сегодня
        elif date_filter == 'week':
            start_date = сегодня - timezone.timedelta(days=7)
            end_date = сегодня
        elif date_filter == 'month':
            start_date = сегодня.replace(day=1)
            end_date = сегодня
        elif date_filter == 'year':
            start_date = сегодня.replace(month=1, day=1)
            end_date = сегодня
    
    if start_date_str:
        date_filter_applied = True
        start_date = timezone.datetime.strptime(start_date_str, '%Y-%m-%d').date()
    if end_date_str:
        date_filter_applied = True
        end_date = timezone.datetime.strptime(end_date_str, '%Y-%m-%d').date()
    
    # Применяем фильтрацию по дате только если она указана
    if date_filter_applied and start_date and end_date:
        выплаты_фильтрованные = выплаты.filter(
            дата_создания__date__range=[start_date, end_date]
        )
    else:
        # Если фильтр по дате не применен, показываем все выплаты
        выплаты_фильтрованные = выплаты
        # Для расчета ожидаемой суммы используем текущий месяц по умолчанию
        start_date = сегодня.replace(day=1)
        end_date = сегодня
    
    # Применяем фильтрацию по организациям
    выбранные_организации = []
    if организации_ids:
        выбранные_организации = Организация.objects.filter(id__in=организации_ids)
        # Находим выплаты для выбранных организаций
        выплаты_фильтрованные = выплаты_фильтрованные.filter(
            полное_наименование__in=[org.ПолноеНаименование for org in выбранные_организации]
        )
    
    # Рассчитываем ожидаемую сумму для периода
    ожидаемая_сумма_периода = 0
    years_in_period = set()
    
    # Собираем все годы в периоде
    current_date = start_date
    while current_date <= end_date:
        years_in_period.add(current_date.year)
        current_date += timezone.timedelta(days=1)
    
    # Суммируем ожидаемые суммы по всем годам в периоде
    for year in years_in_period:
        try:
            годовая_сумма = ОжидаемаяСуммаВыплат.objects.get(год=year)
            # Рассчитываем часть года, попадающую в период
            period_start = max(start_date, timezone.datetime(year, 1, 1).date())
            period_end = min(end_date, timezone.datetime(year, 12, 31).date())
            
            if period_start <= period_end:
                days_in_period = (period_end - period_start).days + 1
                days_in_year = 366 if годовая_сумма.is_leap_year() else 365
                
                daily_rate = годовая_сумма.ожидаемая_сумма_год / days_in_year
                ожидаемая_сумма_периода += daily_rate * days_in_period
                
        except ОжидаемаяСуммаВыплат.DoesNotExist:
            continue
    
    # Фактические выплаты за период (только выплаченные для сравнения с ожидаемой суммой)
    выплаты_выплаченные = выплаты_фильтрованные.filter(статус='выплачено')
    сумма_выплат_периода = выплаты_выплаченные.aggregate(total=models.Sum('сумма'))['total'] or 0
    
    # Подсчет итогов и статистики (по всем выплатам в фильтре)
    total_amount = выплаты_фильтрованные.aggregate(total=models.Sum('сумма'))['total'] or 0
    
    status_counts = {
        'ожидает': выплаты_фильтрованные.filter(статус='ожидает').count(),
        'выплачено': выплаты_фильтрованные.filter(статус='выплачено').count(),
        'отменено': выплаты_фильтрованные.filter(статус='отменено').count(),
    }
    
    status_amounts = {
        'ожидает': выплаты_фильтрованные.filter(статус='ожидает').aggregate(total=models.Sum('сумма'))['total'] or 0,
        'выплачено': выплаты_фильтрованные.filter(статус='выплачено').aggregate(total=models.Sum('сумма'))['total'] or 0,
        'отменено': выплаты_фильтрованные.filter(статус='отменено').aggregate(total=models.Sum('сумма'))['total'] or 0,
    }
    
    context = {
        'выплаты': выплаты_фильтрованные,
        'все_организации': все_организации,
        'выбранные_организации': [str(org.id) for org in выбранные_организации],
        'total_amount': total_amount,
        'status_counts': status_counts,
        'status_amounts': status_amounts,
        'date_filter': date_filter,
        'start_date': start_date_str,
        'end_date': end_date_str,
        'организации_ids': организации_ids,
        # Новые переменные для ожидаемых сумм
        'ожидаемая_сумма_периода': ожидаемая_сумма_периода,
        'сумма_выплат_периода': сумма_выплат_периода,
        'разница_периода': ожидаемая_сумма_периода - сумма_выплат_периода,
        'period_start_date': start_date,
        'period_end_date': end_date,
        'date_filter_applied': date_filter_applied,
    }
    return render(request, 'index/выплаты_list.html', context)

@custom_user_required
def выплата_update_status(request, pk):
    """Обновление статуса выплаты"""
    выплата = get_object_or_404(Выплаты, pk=pk)
    
    # Проверяем, не заблокирована ли выплата
    if выплата.заблокировано:
        messages.error(request, 'Эта выплата уже выполнена и не может быть изменена.')
        return redirect('выплаты_list')
    
    if request.method == 'POST':
        новый_статус = request.POST.get('статус')
        if новый_статус in dict(Выплаты.СТАТУС_ВЫБОР):
            выплата.статус = новый_статус
            выплата.save()
            messages.success(request, f'Статус выплаты обновлен на "{dict(Выплаты.СТАТУС_ВЫБОР)[новый_статус]}"')
        return redirect('выплаты_list')
    
    return redirect('выплаты_list')

def generate_payment_document(выплата):
    """Генерация Word документа для выплаты"""
    doc = Document()
    
    # Заголовок учреждения
    title = doc.add_paragraph()
    title_run = title.add_run('Муниципальное казенное учреждение «Финансово-бюджетная палата»\n')
    title_run.bold = True
    title_run.font.size = Inches(0.14)
    title.add_run('муниципального образования «Лениногорский муниципальный район» Республики Татарстан\n\n')
    
    # Название документа
    doc_title = doc.add_paragraph()
    doc_title.alignment = 1  # Центрирование
    doc_title_run = doc_title.add_run('Уведомление об изменении бюджета\n\n')
    doc_title_run.bold = True
    doc_title_run.font.size = Inches(0.16)
    
    # Информация о выплате
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run(f'Код выплаты: {выплата.код_выплаты}\n')
    info_paragraph.add_run(f'Наименование организации: {выплата.полное_наименование}\n')
    info_paragraph.add_run(f'ИНН: {выплата.инн}\n')
    info_paragraph.add_run(f'Счет получателя: {выплата.счет_получателя}\n')
    info_paragraph.add_run(f'Вид выплаты: {выплата.вид}\n')
    info_paragraph.add_run(f'Сумма выплаты: {выплата.сумма} руб.\n')
    info_paragraph.add_run(f'Дата выполнения: {выплата.дата_выполнения.strftime("%d.%m.%Y") if выплата.дата_выполнения else "Не указана"}\n')
    info_paragraph.add_run(f'Назначение платежа: {выплата.назначение_платежа}\n\n')
    
    # Подпись
    signature = doc.add_paragraph('\n\n\n')
    signature.add_run('_________________________/_________________________')
    signature.alignment = 2  # Выравнивание по правому краю
    signature.add_run('\nПодпись\t\t\t\t\t\tФИО')
    
    # Сохраняем документ в памяти
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def download_payment_document(request, pk):
    """Скачивание документа выплаты"""
    выплата = get_object_or_404(Выплаты, pk=pk)
    
    # Проверяем, можно ли создавать документ (только для выплаченных)
    if выплата.статус != 'выплачено':
        messages.error(request, 'Документ можно скачать только для выплаченных заявок')
        return redirect('выплаты_list')
    
    file_stream = generate_payment_document(выплата)
    
    response = HttpResponse(
        file_stream.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Уведомление_о_выплате_{выплата.код_выплаты}.docx"'
    
    return response

def download_organization_payment_document(request, pk):
    """Скачивание документа выплаты для организации"""
    заявка = get_object_or_404(Заявка, pk=pk)
    
    # Проверяем, есть ли связанная выплата со статусом "выплачено"
    try:
        выплата = Выплаты.objects.get(заявка=заявка, статус='выплачено')
    except Выплаты.DoesNotExist:
        messages.error(request, 'Документ доступен только для выплаченных заявок')
        return redirect('organization_applications')
    
    file_stream = generate_payment_document(выплата)
    
    response = HttpResponse(
        file_stream.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Уведомление_о_выплате_{выплата.код_выплаты}.docx"'
    
    return response

@custom_user_required
def ожидаемые_суммы_list(request):
    """Список ожидаемых сумм по годам"""
    ожидаемые_суммы = ОжидаемаяСуммаВыплат.objects.all().order_by('-год')
    
    # Рассчитываем статистику за текущий год
    текущий_год = timezone.now().year
    текущий_месяц = timezone.now().month
    
    try:
        текущая_ожидаемая_сумма = ОжидаемаяСуммаВыплат.objects.get(год=текущий_год)
        # Ожидаемая сумма за текущий месяц
        ожидаемая_сумма_текущий_месяц = текущая_ожидаемая_сумма.get_expected_for_month(текущий_месяц)
    except ОжидаемаяСуммаВыплат.DoesNotExist:
        текущая_ожидаемая_сумма = None
        ожидаемая_сумма_текущий_месяц = 0
    
    # Фактические выплаты за текущий месяц
    выплаты_текущий_месяц = Выплаты.objects.filter(
        дата_создания__year=текущий_год,
        дата_создания__month=текущий_месяц
    )
    сумма_выплат_текущий_месяц = выплаты_текущий_месяц.aggregate(total=models.Sum('сумма'))['total'] or 0
    
    # Рассчитываем разницу (ожидаемая - фактическая)
    if ожидаемая_сумма_текущий_месяц:
        разница = ожидаемая_сумма_текущий_месяц - сумма_выплат_текущий_месяц
    else:
        разница = None
    
    context = {
        'ожидаемые_суммы': ожидаемые_суммы,
        'текущая_ожидаемая_сумма': текущая_ожидаемая_сумма,
        'ожидаемая_сумма_текущий_месяц': ожидаемая_сумма_текущий_месяц,
        'сумма_выплат_текущий_месяц': сумма_выплат_текущий_месяц,
        'разница': разница,
        'текущий_год': текущий_год,
        'текущий_месяц': текущий_месяц,
    }
    return render(request, 'index/ожидаемые_суммы_list.html', context)

@custom_user_required
def ожидаемая_сумма_create(request):
    """Создание ожидаемой суммы"""
    if request.method == 'POST':
        form = ОжидаемаяСуммаФорма(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Ожидаемая сумма успешно добавлена')
            return redirect('ожидаемые_суммы_list')
    else:
        # Устанавливаем текущий год по умолчанию
        текущий_год = timezone.now().year
        form = ОжидаемаяСуммаФорма(initial={'год': текущий_год})
    
    return render(request, 'index/ожидаемая_сумма_form.html', {'form': form})

@custom_user_required
def ожидаемая_сумма_update(request, pk):
    """Редактирование ожидаемой суммы"""
    ожидаемая_сумма = get_object_or_404(ОжидаемаяСуммаВыплат, pk=pk)
    
    if request.method == 'POST':
        form = ОжидаемаяСуммаФорма(request.POST, instance=ожидаемая_сумма)
        if form.is_valid():
            form.save()
            messages.success(request, 'Ожидаемая сумма успешно обновлена')
            return redirect('ожидаемые_суммы_list')
    else:
        form = ОжидаемаяСуммаФорма(instance=ожидаемая_сумма)
    
    return render(request, 'index/ожидаемая_сумма_form.html', {'form': form})

@custom_user_required
def ожидаемая_сумма_delete(request, pk):
    """Удаление ожидаемой суммы"""
    ожидаемая_сумма = get_object_or_404(ОжидаемаяСуммаВыплат, pk=pk)
    
    if request.method == 'POST':
        ожидаемая_сумма.delete()
        messages.success(request, 'Ожидаемая сумма удалена')
        return redirect('ожидаемые_суммы_list')
    
    return render(request, 'index/ожидаемая_сумма_confirm_delete.html', {'ожидаемая_сумма': ожидаемая_сумма})

# ==================== ОТЧЕТЫ ====================
@custom_user_required
def отчеты_list(request):
    """Список всех отчетов"""
    отчеты = ОтчетВыплат.objects.all()
    return render(request, 'index/отчеты_list.html', {'отчеты': отчеты})

@custom_user_required
def отчет_create_from_payments(request):
    """Создание отчета на основе текущих фильтров выплат"""
    if request.method == 'POST':
        form = ОтчетФорма(request.POST)
        if form.is_valid():
            отчет = form.save(commit=False)
            
            # Получаем параметры фильтрации из сессии или запроса
            date_filter = request.session.get('current_date_filter', '')
            start_date = request.session.get('current_start_date', '')
            end_date = request.session.get('current_end_date', '')
            организации_ids = request.session.get('current_организации_ids', [])
            
            # Если даты не указаны в форме, используем из фильтров
            if not отчет.дата_начала and start_date:
                отчет.дата_начала = datetime.strptime(start_date, '%Y-%m-%d').date()
            if not отчет.дата_окончания and end_date:
                отчет.дата_окончания = datetime.strptime(end_date, '%Y-%m-%d').date()
            
            # Если даты все еще не установлены, используем текущий месяц
            if not отчет.дата_начала:
                отчет.дата_начала = timezone.now().replace(day=1).date()
            if not отчет.дата_окончания:
                отчет.дата_окончания = timezone.now().date()
            
            # Получаем выплаты за период
            выплаты = Выплаты.objects.filter(
                дата_создания__date__range=[отчет.дата_начала, отчет.дата_окончания],
                статус='выплачено'
            )
            
            # Применяем фильтр по организациям если они выбраны
            if организации_ids:
                выбранные_организации = Организация.objects.filter(id__in=организации_ids)
                выплаты = выплаты.filter(
                    полное_наименование__in=[org.ПолноеНаименование for org in выбранные_организации]
                )
            
            отчет.общая_сумма = выплаты.aggregate(total=Sum('сумма'))['total'] or 0
            отчет.количество_выплат = выплаты.count()
            отчет.save()
            
            # Сохраняем выбранные организации
            if организации_ids:
                отчет.организации.set(Организация.objects.filter(id__in=организации_ids))
            
            messages.success(request, 'Отчет успешно создан')
            return redirect('отчеты_list')
    else:
        # Автозаполнение дат из текущих фильтров
        initial = {}
        date_filter = request.GET.get('date_filter', '')
        start_date = request.GET.get('start_date', '')
        end_date = request.GET.get('end_date', '')
        организации_ids = request.GET.getlist('организации')
        
        # Определяем даты на основе быстрого фильтра
        сегодня = timezone.now().date()
        calculated_start_date = None
        calculated_end_date = None
        
        if date_filter:
            if date_filter == 'today':
                calculated_start_date = сегодня
                calculated_end_date = сегодня
            elif date_filter == 'week':
                calculated_start_date = сегодня - timezone.timedelta(days=7)
                calculated_end_date = сегодня
            elif date_filter == 'month':
                calculated_start_date = сегодня.replace(day=1)
                calculated_end_date = сегодня
            elif date_filter == 'year':
                calculated_start_date = сегодня.replace(month=1, day=1)
                calculated_end_date = сегодня
        
        # Сохраняем текущие фильтры в сессии
        request.session['current_date_filter'] = date_filter
        
        # Устанавливаем даты в initial
        if start_date:
            request.session['current_start_date'] = start_date
            initial['дата_начала'] = start_date
        elif calculated_start_date:
            initial['дата_начала'] = calculated_start_date.strftime('%Y-%m-%d')
        
        if end_date:
            request.session['current_end_date'] = end_date
            initial['дата_окончания'] = end_date
        elif calculated_end_date:
            initial['дата_окончания'] = calculated_end_date.strftime('%Y-%m-%d')
        
        if организации_ids:
            request.session['current_организации_ids'] = организации_ids
            initial['организации'] = организации_ids
        
        # Автозаполнение названия
        название_части = []
        if date_filter:
            if date_filter == 'today':
                название_части.append('сегодня')
            elif date_filter == 'week':
                название_части.append('неделю')
            elif date_filter == 'month':
                название_части.append('месяц')
            elif date_filter == 'year':
                название_части.append('год')
        
        if организации_ids:
            org_count = len(организации_ids)
            if org_count == 1:
                org = Организация.objects.get(id=организации_ids[0])
                название_части.append(f'орг. {org.ПолноеНаименование[:20]}')
            else:
                название_части.append(f'{org_count} орг.')
        
        if название_части:
            initial['название'] = f'Отчет за {" ".join(название_части)}'
        else:
            initial['название'] = f'Отчет за период'
        
        form = ОтчетФорма(initial=initial)
    
    return render(request, 'index/отчет_create.html', {'form': form})

@custom_user_required
def отчет_детали(request, pk):
    """Детали отчета"""
    отчет = get_object_or_404(ОтчетВыплат, pk=pk)
    
    # Получаем выплаты для этого отчета
    выплаты = Выплаты.objects.filter(
        дата_создания__date__range=[отчет.дата_начала, отчет.дата_окончания],
        статус='выплачено'
    )
    
    # Применяем фильтр по организациям если они есть
    if отчет.организации.exists():
        выплаты = выплаты.filter(
            полное_наименование__in=[org.ПолноеНаименование for org in отчет.организации.all()]
        )
    
    выплаты = выплаты.order_by('-дата_создания')
    
    # Данные для диаграммы
    данные_диаграммы = выплаты.values('вид__Название').annotate(
        total=Sum('сумма'),
        count=Count('id')
    ).order_by('-total')
    
    context = {
        'отчет': отчет,
        'выплаты': выплаты,
        'данные_диаграммы': list(данные_диаграммы),
    }
    return render(request, 'index/отчет_детали.html', context)

@custom_user_required
def отчет_download(request, pk):
    """Скачивание отчета в Word (официальный стиль)"""
    отчет = get_object_or_404(ОтчетВыплат, pk=pk)
    
    # Получаем выплаты для отчета
    выплаты = Выплаты.objects.filter(
        дата_создания__date__range=[отчет.дата_начала, отчет.дата_окончания],
        статус='выплачено'
    )
    
    # Применяем фильтр по организациям если они есть
    if отчет.организации.exists():
        выплаты = выплаты.filter(
            полное_наименование__in=[org.ПолноеНаименование for org in отчет.организации.all()]
        )
    
    выплаты = выплаты.order_by('-дата_создания')
    
    # Создаем Word документ
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    
    # Настройка шрифта по умолчанию
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    # Верхний колонтитул (шапка) - центрирование
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    header_run = header_para.add_run('МУНИЦИПАЛЬНОЕ КАЗЕННОЕ УЧРЕЖДЕНИЕ\n')
    header_run.bold = True
    header_run.font.size = Pt(14)
    
    header_run2 = header_para.add_run('«ФИНАНСОВО-БЮДЖЕТНАЯ ПАЛАТА»\n')
    header_run2.bold = True
    header_run2.font.size = Pt(14)
    
    header_para.add_run('муниципального образования «Лениногорский муниципальный район»\n')
    header_para.add_run('Республики Татарстан\n\n')
    
    # Название документа
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('ОТЧЕТ\n')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.add_run('об исполнении бюджета\n\n')
    
    # Информация об отчете
    doc.add_paragraph(f'Название отчета: {отчет.название}')
    doc.add_paragraph(f'Период: с {отчет.дата_начала.strftime("%d.%m.%Y")} по {отчет.дата_окончания.strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Дата формирования: {отчет.дата_создания.strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph('')
    
    # Информация об организациях
    if отчет.организации.exists():
        org_names = ', '.join([org.ПолноеНаименование for org in отчет.организации.all()])
        doc.add_paragraph(f'Организации: {org_names}')
        doc.add_paragraph('')
    
    # КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ
    doc.add_paragraph('1. КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ')
    doc.paragraphs[-1].runs[0].bold = True
    
    # Создаем таблицу для показателей
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_table.cell(0, 0).text = 'Общая сумма выплат'
    info_table.cell(0, 1).text = f'{отчет.общая_сумма:,.2f} рублей'
    info_table.cell(1, 0).text = 'Количество выплат'
    info_table.cell(1, 1).text = f'{отчет.количество_выплат} шт.'
    info_table.cell(2, 0).text = 'Период'
    info_table.cell(2, 1).text = f'{отчет.дата_начала.strftime("%d.%m.%Y")} - {отчет.дата_окончания.strftime("%d.%m.%Y")}'
    info_table.cell(3, 0).text = 'Дата формирования'
    info_table.cell(3, 1).text = f'{отчет.дата_создания.strftime("%d.%m.%Y %H:%M")}'
    
    # Жирный шрифт для первого столбца
    for row in info_table.rows:
        for paragraph in row.cells[0].paragraphs:
            if paragraph.runs:
                paragraph.runs[0].bold = True
            else:
                run = paragraph.add_run(row.cells[0].text)
                run.bold = True
                paragraph.clear()
                paragraph.add_run(row.cells[0].text).bold = True
    
    doc.add_paragraph('')
    
    # Таблица выплат
    doc.add_paragraph('2. ПЕРЕЧЕНЬ ВЫПЛАТ')
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph('')
    
    if выплаты.exists():
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        headers = ['№ п/п', 'Код выплаты', 'Наименование организации', 'Вид выплаты', 'Сумма (руб.)', 'Дата выплаты']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].bold = True
                else:
                    run = paragraph.add_run(header)
                    run.bold = True
        
        # Заполняем таблицу данными
        for idx, выплата in enumerate(выплаты, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = выплата.код_выплаты
            row.cells[2].text = выплата.полное_наименование[:50] if len(выплата.полное_наименование) > 50 else выплата.полное_наименование
            row.cells[3].text = str(выплата.вид) if выплата.вид else '-'
            row.cells[4].text = f'{выплата.сумма:,.2f}'
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[5].text = выплата.дата_создания.strftime('%d.%m.%Y')
            row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем итоговую строку
        row = table.add_row()
        row.cells[0].text = ''
        row.cells[1].text = ''
        row.cells[2].text = ''
        row.cells[3].text = 'ИТОГО:'
        row.cells[4].text = f'{отчет.общая_сумма:,.2f}'
        row.cells[5].text = ''
        
        # Выделяем итоговую строку жирным
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].bold = True
                elif paragraph.text:
                    run = paragraph.add_run(paragraph.text)
                    run.bold = True
        
        doc.add_paragraph('')
        doc.add_paragraph(f'Общая сумма выплат: {отчет.общая_сумма:,.2f} рублей.')
        
    else:
        doc.add_paragraph('Выплаты, соответствующие критериям, не найдены.')
        doc.paragraphs[-1].runs[0].bold = True
    
    doc.add_paragraph('')
    
    # Заключение
    doc.add_paragraph('3. ЗАКЛЮЧЕНИЕ')
    doc.paragraphs[-1].runs[0].bold = True
    
    conclusion_para = doc.add_paragraph()
    conclusion_para.add_run(f'За отчетный период с {отчет.дата_начала.strftime("%d.%m.%Y")} по {отчет.дата_окончания.strftime("%d.%m.%Y")} было произведено {отчет.количество_выплат} выплат на общую сумму {отчет.общая_сумма:,.2f} рублей. ')
    
    if отчет.организации.exists():
        conclusion_para.add_run(f'Отчет сформирован для следующих организаций: ')
        orgs = ', '.join([org.ПолноеНаименование for org in отчет.организации.all()])
        conclusion_para.add_run(f'{orgs}. ')
    
    conclusion_para.add_run('Все выплаты соответствуют утвержденному плану и произведены в установленные сроки.')
    
    # Подписи
    doc.add_paragraph('\n\n')
    doc.add_paragraph('4. ПОДПИСИ')
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph('')
    
    # Получаем текущего пользователя из сессии
    from .models import Пользователь
    user_id = request.session.get('user_id')
    user = None
    if user_id:
        try:
            user = Пользователь.objects.get(pk=user_id)
        except:
            pass
    
    if user:
        doc.add_paragraph('Ответственный сотрудник:')
        doc.add_paragraph(f'{user.Фамилия} {user.Имя} {user.Отчество or ""}')
    else:
        doc.add_paragraph('Ответственный сотрудник:')
        doc.add_paragraph('_________________________')
    
    doc.add_paragraph('__________________________')
    doc.add_paragraph('(подпись)')
    
    doc.add_paragraph('')
    doc.add_paragraph('Место печати')
    doc.add_paragraph('М.П.')
    
    # Нижний колонтитул
    doc.add_paragraph('')
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(f'Документ сформирован автоматически {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    for run in footer_para.runs:
        run.font.size = Pt(10)
    
    # Сохраняем документ
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    response = HttpResponse(
        file_stream.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Отчет_{отчет.название}_{datetime.now().strftime("%Y%m%d_%H%M")}.docx"'
    
    return response

@custom_user_required
def отчет_chart(request, pk):
    """Страница с диаграммой отчета"""
    отчет = get_object_or_404(ОтчетВыплат, pk=pk)
    
    # Получаем выплаты для этого отчета
    выплаты = Выплаты.objects.filter(
        дата_создания__date__range=[отчет.дата_начала, отчет.дата_окончания],
        статус='выплачено'
    )
    
    # Применяем фильтр по организациям если они есть
    if отчет.организации.exists():
        выплаты = выплаты.filter(
            полное_наименование__in=[org.ПолноеНаименование for org in отчет.организации.all()]
        )
    
    # Данные для диаграммы по видам выплат
    данные_диаграммы = выплаты.values('вид__Название').annotate(
        total=Sum('сумма'),
        count=Count('id')
    ).order_by('-total')
    
    # Данные для диаграммы по организациям
    данные_организаций = выплаты.values('полное_наименование').annotate(
        total=Sum('сумма'),
        count=Count('id')
    ).order_by('-total')[:10]  # Топ 10 организаций
    
    # Отладочная информация
    print(f"Отчет: {отчет.название}")
    print(f"Количество выплат: {выплаты.count()}")
    print(f"Данные диаграммы: {list(данные_диаграммы)}")
    print(f"Данные организаций: {list(данные_организаций)}")
    
    context = {
        'отчет': отчет,
        'выплаты': выплаты,
        'данные_диаграммы': list(данные_диаграммы),
        'данные_организаций': list(данные_организаций),
    }
    return render(request, 'index/отчет_chart.html', context)

@custom_user_required
def planning_payments(request):
    """Страница выплат для специалиста по планированию"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Отдел по планированию и анализу':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    # Автоматически создаем выплаты из принятых заявок
    create_payments_from_approved_applications()
    
    # ТОЛЬКО выплаченные выплаты (статус 'выплачено')
    выплаты = Выплаты.objects.filter(статус='выплачено').order_by('-дата_создания')
    
    # Получаем все организации для фильтра
    все_организации = Организация.objects.all()
    
    # Получаем параметры фильтрации
    date_filter = request.GET.get('date_filter', '')
    start_date_str = request.GET.get('start_date', '')
    end_date_str = request.GET.get('end_date', '')
    организации_ids = request.GET.getlist('организации')
    
    # Переменные для выбранных организаций
    выбранные_организации = []
    
    # Применяем фильтрацию по дате
    if date_filter:
        сегодня = timezone.now().date()
        if date_filter == 'today':
            выплаты = выплаты.filter(дата_создания__date=сегодня)
        elif date_filter == 'week':
            неделя_назад = сегодня - timezone.timedelta(days=7)
            выплаты = выплаты.filter(дата_создания__date__range=[неделя_назад, сегодня])
        elif date_filter == 'month':
            начало_месяца = сегодня.replace(day=1)
            выплаты = выплаты.filter(дата_создания__date__range=[начало_месяца, сегодня])
        elif date_filter == 'year':
            начало_года = сегодня.replace(month=1, day=1)
            выплаты = выплаты.filter(дата_создания__date__range=[начало_года, сегодня])
    
    if start_date_str and end_date_str:
        start_date = timezone.datetime.strptime(start_date_str, '%Y-%m-%d').date()
        end_date = timezone.datetime.strptime(end_date_str, '%Y-%m-%d').date()
        выплаты = выплаты.filter(дата_создания__date__range=[start_date, end_date])
    
    # Применяем фильтрацию по организациям
    if организации_ids:
        выбранные_организации = Организация.objects.filter(id__in=организации_ids)
        выплаты = выплаты.filter(
            полное_наименование__in=[org.ПолноеНаименование for org in выбранные_организации]
        )
    
    # Статистика
    total_amount = выплаты.aggregate(total=Sum('сумма'))['total'] or 0
    total_count = выплаты.count()
    
    # Рассчитываем средний чек
    if total_count > 0:
        average_amount = total_amount / total_count
    else:
        average_amount = 0
    
    # Ожидаемая сумма для периода
    сегодня = timezone.now().date()
    
    # Определяем даты периода для расчета ожидаемой суммы
    period_start_date = сегодня.replace(day=1)  # по умолчанию начало месяца
    period_end_date = сегодня  # по умолчанию сегодня
    
    if start_date_str:
        period_start_date = timezone.datetime.strptime(start_date_str, '%Y-%m-%d').date()
    if end_date_str:
        period_end_date = timezone.datetime.strptime(end_date_str, '%Y-%m-%d').date()
    
    # Если используется быстрый фильтр по дате
    if date_filter and not start_date_str:
        if date_filter == 'today':
            period_start_date = сегодня
            period_end_date = сегодня
        elif date_filter == 'week':
            period_start_date = сегодня - timezone.timedelta(days=7)
            period_end_date = сегодня
        elif date_filter == 'month':
            period_start_date = сегодня.replace(day=1)
            period_end_date = сегодня
        elif date_filter == 'year':
            period_start_date = сегодня.replace(month=1, day=1)
            period_end_date = сегодня
    
    период_в_днях = (period_end_date - period_start_date).days + 1
    год = period_start_date.year
    
    try:
        ожидаемая_сумма = ОжидаемаяСуммаВыплат.objects.get(год=год)
        days_in_year = 366 if ожидаемая_сумма.is_leap_year() else 365
        daily_rate = ожидаемая_сумма.ожидаемая_сумма_год / days_in_year
        ожидаемая_сумма_периода = daily_rate * период_в_днях
    except ОжидаемаяСуммаВыплат.DoesNotExist:
        ожидаемая_сумма_периода = 0
    
    context = {
        'выплаты': выплаты,
        'все_организации': все_организации,
        'выбранные_организации': [str(org.id) for org in выбранные_организации],
        'total_amount': total_amount,
        'total_count': total_count,
        'average_amount': average_amount,  # ДОБАВЛЕНО
        'date_filter': date_filter,
        'start_date': start_date_str,
        'end_date': end_date_str,
        'организации_ids': организации_ids,
        'ожидаемая_сумма_периода': ожидаемая_сумма_периода,
        'user': user,
        'period_start_date': period_start_date,
        'period_end_date': period_end_date,
    }
    return render(request, 'index/planning_payments.html', context)

# ==================== ЛИЧНЫЙ КАБИНЕТ СПЕЦИАЛИСТА ПО ПЛАНИРОВАНИЮ ====================
@custom_user_required
def planning_dashboard(request):
    """Личный кабинет специалиста по планированию и анализу"""
    # Проверяем сессию
    if not request.session.get('user_id') or request.session.get('user_role') != 'planning':
        return redirect('login')
    
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        # Дополнительная проверка отдела
        if not user.Отдел or user.Отдел.Название != 'Отдел по планированию и анализу':
            request.session.flush()
            return redirect('login')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    # Получаем статистику для дашборда
    текущий_год = timezone.now().year
    текущий_месяц = timezone.now().month
    
    # Статистика выплат
    выплаты_месяц = Выплаты.objects.filter(
        дата_создания__year=текущий_год,
        дата_создания__month=текущий_месяц,
        статус='выплачено'
    )
    
    сумма_выплат_месяц = выплаты_месяц.aggregate(total=Sum('сумма'))['total'] or 0
    количество_выплат_месяц = выплаты_месяц.count()
    
    # Ожидаемая сумма
    try:
        ожидаемая_сумма = ОжидаемаяСуммаВыплат.objects.get(год=текущий_год)
        ожидаемая_сумма_месяц = ожидаемая_сумма.get_expected_for_month(текущий_месяц)
    except ОжидаемаяСуммаВыплат.DoesNotExist:
        ожидаемая_сумма_месяц = 0
    
    # ВАЖНО: Отклонение = ПЛАН - ФАКТ (а не факт - план)
    отклонение = ожидаемая_сумма_месяц - сумма_выплат_месяц
    
    # Отчеты
    отчеты = ОтчетПланировщика.objects.all().order_by('-дата_создания')[:5]
    всего_отчетов = ОтчетПланировщика.objects.count()
    
    # Месяц и год для отображения
    месяцы = {
        1: 'Январь', 2: 'Февраль', 3: 'Март', 4: 'Апрель',
        5: 'Май', 6: 'Июнь', 7: 'Июль', 8: 'Август',
        9: 'Сентябрь', 10: 'Октябрь', 11: 'Ноябрь', 12: 'Декабрь'
    }
    название_месяца = месяцы.get(текущий_месяц, '')
    
    context = {
        'user': user,
        'сумма_выплат_месяц': сумма_выплат_месяц,
        'количество_выплат_месяц': количество_выплат_месяц,
        'ожидаемая_сумма_месяц': ожидаемая_сумма_месяц,
        'отклонение': отклонение,  # Теперь это план - факт
        'отчеты': отчеты,
        'всего_отчетов': всего_отчетов,
        'текущий_год': текущий_год,
        'текущий_месяц': текущий_месяц,
        'название_месяца': название_месяца,
    }
    
    print(f"→ Отклонение (план-факт): {отклонение}")
    return render(request, 'index/planning_dashboard.html', context)
@custom_user_required
def planning_reports_list(request):
    """Список отчетов для специалиста по планированию"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Отдел по планированию и анализу':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    отчеты = ОтчетВыплат.objects.all().order_by('-дата_создания')
    
    context = {
        'отчеты': отчеты,
        'user': user,
    }
    return render(request, 'index/planning_reports_list.html', context)

@custom_user_required
def planning_create_report(request):
    """Создание отчета для специалиста по планированию"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Отдел по планированию и анализу':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    if request.method == 'POST':
        form = ОтчетФорма(request.POST)
        if form.is_valid():
            отчет = form.save(commit=False)
            
            # Используем select_related и только необходимые поля
            выплаты = Выплаты.objects.filter(
                дата_создания__date__range=[отчет.дата_начала, отчет.дата_окончания],
                статус='выплачено'
            ).select_related('вид')  # Явная загрузка связанных данных
            
            # Применяем фильтр по организациям если они выбраны
            if form.cleaned_data.get('организации'):
                выбранные_организации = form.cleaned_data['организации']
                выплаты = выплаты.filter(
                    полное_наименование__in=[org.ПолноеНаименование for org in выбранные_организации]
                )
            
            # Вычисляем агрегаты БЕЗ обращения к связанным объектам
            агрегаты = выплаты.aggregate(
                total=Sum('сумма'),
                count=Count('id')
            )
            
            отчет.общая_сумма = агрегаты['total'] or 0
            отчет.количество_выплат = агрегаты['count'] or 0
            отчет.save()
            
            # Сохраняем выбранные организации
            if form.cleaned_data.get('организации'):
                отчет.организации.set(form.cleaned_data['организации'])
            
            messages.success(request, 'Отчет успешно создан')
            return redirect('planning_reports_list')
    else:
        form = ОтчетФорма()
    
    return render(request, 'index/planning_create_report.html', {'form': form, 'user': user})

@custom_user_required
def planning_report_analysis(request, pk):
    """Анализ отчета с графиками для специалиста по планированию"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Отдел по планированию и анализу':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    отчет = get_object_or_404(ОтчетВыплат, pk=pk)
    
    # Получаем выплаты для отчета
    выплаты = Выплаты.objects.filter(
        дата_создания__date__range=[отчет.дата_начала, отчет.дата_окончания],
        статус='выплачено'
    )
    
    # Применяем фильтр по организациям если они есть
    if отчет.организации.exists():
        выплаты = выплаты.filter(
            полное_наименование__in=[org.ПолноеНаименование for org in отчет.организации.all()]
        )
    
    # Данные для анализа плана vs факта
    период_в_днях = (отчет.дата_окончания - отчет.дата_начала).days + 1
    год = отчет.дата_начала.year
    
    try:
        ожидаемая_сумма = ОжидаемаяСуммаВыплат.objects.get(год=год)
        # Рассчитываем ожидаемую сумму за период
        days_in_year = 366 if ожидаемая_сумма.is_leap_year() else 365
        daily_rate = ожидаемая_сумма.ожидаемая_сумма_год / days_in_year
        ожидаемая_сумма_периода = daily_rate * период_в_днях
    except ОжидаемаяСуммаВыплат.DoesNotExist:
        ожидаемая_сумма_периода = 0
    
    фактическая_сумма = выплаты.aggregate(total=Sum('сумма'))['total'] or 0
    отклонение = фактическая_сумма - ожидаемая_сумма_периода
    
    # Данные для диаграммы по видам выплат
    данные_диаграммы = выплаты.values('вид__Название').annotate(
        total=Sum('сумма'),
        count=Count('id')
    ).order_by('-total')
    
    # Данные для диаграммы по месяцам
    monthly_data = []
    current_date = отчет.дата_начала.replace(day=1)
    while current_date <= отчет.дата_окончания:
        month_start = current_date
        month_end = (current_date.replace(day=28) + timezone.timedelta(days=4)).replace(day=1) - timezone.timedelta(days=1)
        if month_end > отчет.дата_окончания:
            month_end = отчет.дата_окончания
        
        month_payments = выплаты.filter(
            дата_создания__date__range=[month_start, month_end]
        )
        month_total = month_payments.aggregate(total=Sum('сумма'))['total'] or 0
        
        # Ожидаемая сумма за месяц
        month_days = (month_end - month_start).days + 1
        month_expected = daily_rate * month_days if ожидаемая_сумма_периода > 0 else 0
        
        monthly_data.append({
            'month': current_date.strftime('%b %Y'),
            'факт': float(month_total),
            'план': float(month_expected),
            'отклонение': float(month_total - month_expected)
        })
        
        current_date = (current_date.replace(day=28) + timezone.timedelta(days=4)).replace(day=1)
    
    # Данные для круговой диаграммы по организациям
    данные_организаций = выплаты.values('полное_наименование').annotate(
        total=Sum('сумма')
    ).order_by('-total')[:10]  # Топ 10 организаций
    
    # Подготовка данных для JavaScript
    chart_data_types = {
        'labels': [item['вид__Название'] or 'Не указан' for item in данные_диаграммы],
        'data': [float(item['total']) for item in данные_диаграммы],
        'counts': [item['count'] for item in данные_диаграммы]
    }
    
    chart_data_organizations = {
        'labels': [item['полное_наименование'] for item in данные_организаций],
        'data': [float(item['total']) for item in данные_организаций]
    }
    
    chart_data_monthly = {
        'labels': [item['month'] for item in monthly_data],
        'факт': [item['факт'] for item in monthly_data],
        'план': [item['план'] for item in monthly_data],
        'отклонение': [item['отклонение'] for item in monthly_data]
    }
    
    context = {
        'отчет': отчет,
        'user': user,
        'выплаты': выплаты,
        'ожидаемая_сумма_периода': ожидаемая_сумма_периода,
        'фактическая_сумма': фактическая_сумма,
        'отклонение': отклонение,
        'monthly_data': monthly_data,
        'процент_выполнения': (фактическая_сумма / ожидаемая_сумма_периода * 100) if ожидаемая_сумма_периода > 0 else 0,
        'chart_data_types': chart_data_types,
        'chart_data_organizations': chart_data_organizations,
        'chart_data_monthly': chart_data_monthly,
    }
    return render(request, 'index/planning_report_analysis.html', context)

def create_payments_from_approved_applications():
    """Автоматически создает выплаты из принятых заявок"""
    # ИСПОЛЬЗУЙТЕ НОВЫЙ related_name
    approved_applications = Заявка.objects.filter(
        статус='approved', 
        связанные_выплаты__isnull=True  # ИЗМЕНИТЕ здесь
    )
    
    for application in approved_applications:
        # Используем get_or_create для избежания дубликатов
        выплата, created = Выплаты.objects.get_or_create(
            заявка=application,
            defaults={
                'полное_наименование': application.полное_наименование,
                'инн': application.инн,
                'счет_получателя': application.счет_получателя,
                'сумма': application.запрашиваемая_сумма,
                'вид': application.вид_выплаты,
                'назначение_платежа': f"Выплата по заявке #{application.id}",
                'статус': 'выплачено'  # Сразу устанавливаем статус "выплачено"
            }
        )
        
        if created:
            print(f"✅ Создана выплата для заявки #{application.id}")
            

# ==================== ОТЧЕТЫ ДЛЯ СПЕЦИАЛИСТА ПО ПЛАНИРОВАНИЮ ====================
@custom_user_required
def planning_reports_list(request):
    """Список отчетов планировщика - СОВСЕМ ДРУГИЕ ОТЧЕТЫ!"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Отдел по планированию и анализу':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    отчеты = ОтчетПланировщика.objects.all().order_by('-дата_создания')
    
    context = {
        'отчеты': отчеты,
        'user': user,
    }
    return render(request, 'index/planning_reports_list.html', context)

@custom_user_required
def planning_create_report(request):
    """Создание отчета планировщика - С ФОКУСОМ НА ОЖИДАЕМЫХ И ФАКТИЧЕСКИХ СУММАХ"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Отдел по планированию и анализу':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    if request.method == 'POST':
        form = ОтчетПланировщикаФорма(request.POST)
        if form.is_valid():
            отчет = form.save(commit=False)
            
            # РАСЧЕТ ОЖИДАЕМОЙ СУММЫ для периода
            период_в_днях = (отчет.дата_окончания - отчет.дата_начала).days + 1
            год = отчет.дата_начала.year
            
            try:
                ожидаемая_сумма_год = ОжидаемаяСуммаВыплат.objects.get(год=год)
                days_in_year = 366 if ожидаемая_сумма_год.is_leap_year() else 365
                daily_rate = ожидаемая_сумма_год.ожидаемая_сумма_год / days_in_year
                отчет.ожидаемая_сумма_периода = daily_rate * период_в_днях
            except ОжидаемаяСуммаВыплат.DoesNotExist:
                отчет.ожидаемая_сумма_периода = 0
            
            # РАСЧЕТ ФАКТИЧЕСКОЙ СУММЫ для периода
            выплаты = Выплаты.objects.filter(
                дата_создания__date__range=[отчет.дата_начала, отчет.дата_окончания],
                статус='выплачено'
            )
            отчет.фактическая_сумма_периода = выплаты.aggregate(total=Sum('сумма'))['total'] or 0
            отчет.количество_выплат = выплаты.count()
            
            # СОБИРАЕМ ДАННЫЕ ДЛЯ ГРАФИКА В WORD
            # Данные по месяцам для графика
            monthly_data = []
            current_date = отчет.дата_начала.replace(day=1)
            
            while current_date <= отчет.дата_окончания:
                month_start = current_date
                month_end = (current_date.replace(day=28) + timezone.timedelta(days=4)).replace(day=1) - timezone.timedelta(days=1)
                if month_end > отчет.дата_окончания:
                    month_end = отчет.дата_окончания
                
                # Фактические выплаты за месяц
                month_payments = выплаты.filter(дата_создания__date__range=[month_start, month_end])
                month_fact = month_payments.aggregate(total=Sum('сумма'))['total'] or 0
                
                # Ожидаемые выплаты за месяц
                month_days = (month_end - month_start).days + 1
                month_expected = daily_rate * month_days if daily_rate > 0 else 0
                
                monthly_data.append({
                    'month': current_date.strftime('%B %Y'),
                    'ожидаемая': float(month_expected),
                    'фактическая': float(month_fact),
                    'отклонение': float(month_fact - month_expected)
                })
                
                current_date = (current_date.replace(day=28) + timezone.timedelta(days=4)).replace(day=1)
            
            # Сохраняем данные для графика
            отчет.данные_графика = {
                'monthly_data': monthly_data,
                'period_days': период_в_днях,
                'year': год,
                'daily_rate': float(daily_rate) if 'daily_rate' in locals() else 0
            }
            
            отчет.save()
            
            messages.success(request, 'Отчет планировщика успешно создан!')
            return redirect('planning_reports_list')
    else:
        form = ОтчетПланировщикаФорма()
    
    return render(request, 'index/planning_create_report.html', {'form': form, 'user': user})

@custom_user_required
def planning_report_detail(request, pk):
    """Детали отчета планировщика"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Отдел по планированию и анализу':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    отчет = get_object_or_404(ОтчетПланировщика, pk=pk)
    
    # ВЫЧИСЛЯЕМ ОТКЛОНЕНИЕ ПРАВИЛЬНО (план - факт)
    отклонение_план_минус_факт = отчет.ожидаемая_сумма_периода - отчет.фактическая_сумма_периода
    
    context = {
        'отчет': отчет,
        'user': user,
        'отклонение_план_минус_факт': отклонение_план_минус_факт,
    }
    return render(request, 'index/planning_report_detail.html', context)

@custom_user_required
def download_planning_report(request, pk):
    """Скачивание отчета планировщика в Word - С ГРАФИКОМ"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    
    отчет = get_object_or_404(ОтчетПланировщика, pk=pk)
    
    doc = Document()
    
    # Настройка шрифта Times New Roman 14pt
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    # Верхний колонтитул
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.add_run('МУНИЦИПАЛЬНОЕ КАЗЕННОЕ УЧРЕЖДЕНИЕ\n').bold = True
    header_para.add_run('«ФИНАНСОВО-БЮДЖЕТНАЯ ПАЛАТА»\n').bold = True
    header_para.add_run('муниципального образования «Лениногорский муниципальный район»\n')
    header_para.add_run('Республики Татарстан\n\n')
    
    # Заголовок
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run('АНАЛИТИЧЕСКИЙ ОТЧЕТ\n').bold = True
    title_para.add_run('о выполнении плановых показателей выплат\n\n')
    
    # Информация об отчете
    doc.add_paragraph(f'Название отчета: {отчет.название}')
    doc.add_paragraph(f'Период анализа: с {отчет.дата_начала.strftime("%d.%m.%Y")} по {отчет.дата_окончания.strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Дата формирования: {отчет.дата_создания.strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph('')
    
    # Ключевые показатели
    doc.add_paragraph('1. КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ').runs[0].bold = True
    
    # Таблица показателей
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    deviation = отчет.ожидаемая_сумма_периода - отчет.фактическая_сумма_периода
    
    table.cell(0, 0).text = 'План (ожидаемая сумма)'
    table.cell(0, 1).text = f'{отчет.ожидаемая_сумма_периода:,.2f} рублей'
    table.cell(1, 0).text = 'Факт (фактическая сумма)'
    table.cell(1, 1).text = f'{отчет.фактическая_сумма_периода:,.2f} рублей'
    table.cell(2, 0).text = 'Отклонение (план - факт)'
    table.cell(2, 1).text = f'{deviation:,.2f} рублей'
    table.cell(3, 0).text = 'Процент выполнения плана'
    table.cell(3, 1).text = f'{отчет.процент_выполнения_плана:.1f}%'
    
    # Жирный шрифт для первого столбца
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph('')
    
    # ГРАФИК - создаем изображение
    try:
        # Данные для графика
        monthly_data = отчет.данные_графика.get('monthly_data', [])
        
        if monthly_data:
            months = [data['month'] for data in monthly_data]
            expected = [data['ожидаемая'] for data in monthly_data]
            actual = [data['фактическая'] for data in monthly_data]
            
            # Создаем график
            plt.figure(figsize=(10, 6))
            x = range(len(months))
            width = 0.35
            
            plt.bar([i - width/2 for i in x], expected, width, label='План (ожидаемые)', color='#3498db', alpha=0.7)
            plt.bar([i + width/2 for i in x], actual, width, label='Факт (фактические)', color='#2ecc71', alpha=0.7)
            
            plt.xlabel('Месяцы')
            plt.ylabel('Сумма (руб.)')
            plt.title('Сравнение плановых и фактических выплат по месяцам')
            plt.xticks(x, months, rotation=45)
            plt.legend()
            plt.tight_layout()
            
            # Сохраняем график в буфер
            buffer = BytesIO()
            plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            
            # Добавляем график в документ
            doc.add_paragraph('2. ГРАФИК: СРАВНЕНИЕ ПЛАНОВЫХ И ФАКТИЧЕСКИХ ВЫПЛАТ').runs[0].bold = True
            doc.add_picture(buffer, width=Inches(6))
            doc.add_paragraph('')
            
    except Exception as e:
        doc.add_paragraph(f'График не может быть отображен: {str(e)}')
        doc.add_paragraph('')
    
    # Аналитический вывод
    doc.add_paragraph('3. АНАЛИТИЧЕСКИЙ ВЫВОД').runs[0].bold = True
    
    if deviation >= 0:
        doc.add_paragraph(f'За анализируемый период достигнута экономия бюджетных средств в размере {deviation:,.2f} рублей. Фактические выплаты оказались ниже плановых показателей.')
    else:
        doc.add_paragraph(f'За анализируемый период допущен перерасход бюджетных средств в размере {abs(deviation):,.2f} рублей. Фактические выплаты превысили плановые показатели.')
    
    doc.add_paragraph(f'План выполнен на {отчет.процент_выполнения_плана:.1f} процентов. За период совершено {отчет.количество_выплат} выплат на общую сумму {отчет.фактическая_сумма_периода:,.2f} рублей.')
    doc.add_paragraph('')
    
    # Рекомендации
    doc.add_paragraph('4. РЕКОМЕНДАЦИИ').runs[0].bold = True
    
    if отчет.процент_выполнения_плана >= 100:
        doc.add_paragraph('• Провести анализ причин перевыполнения плана;')
        doc.add_paragraph('• Рассмотреть возможность корректировки плановых показателей на следующий период;')
        doc.add_paragraph('• Оценить эффективность использования дополнительных средств.')
    elif отчет.процент_выполнения_плана >= 90:
        doc.add_paragraph('• Продолжить мониторинг выполнения плановых показателей;')
        doc.add_paragraph('• Поддерживать текущий уровень эффективности;')
        doc.add_paragraph('• Оптимизировать процессы планирования.')
    else:
        doc.add_paragraph('• Выявить причины недовыполнения плана;')
        doc.add_paragraph('• Разработать план мероприятий по повышению эффективности;')
        doc.add_paragraph('• Усилить контроль за исполнением плановых назначений.')
    
    # Подписи
    doc.add_paragraph('\n\n')
    doc.add_paragraph('5. ПОДПИСИ').runs[0].bold = True
    doc.add_paragraph('')
    doc.add_paragraph('Специалист по планированию и анализу:')
    doc.add_paragraph('_________________________')
    doc.add_paragraph('(подпись)')
    doc.add_paragraph('')
    doc.add_paragraph('М.П.')
    
    # Сохраняем документ
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    response = HttpResponse(
        file_stream.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Аналитический_отчет_{отчет.название}.docx"'
    
    return response

# ==================== СМЕТЫ ЗАТРАТ ОТ АДМИНИСТРАЦИИ ====================

def get_administracia_organization():
    """Вспомогательная функция для получения организации Администрация"""
    try:
        return Организация.objects.get(Логин='Administracia')
    except Организация.DoesNotExist:
        return None


def organization_smeti_list(request):
    """Список смет для организации (Администрация)"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    org_id = request.session.get('organization_id')
    try:
        org = Организация.objects.get(id=org_id)
    except Организация.DoesNotExist:
        return redirect('organization_logout')
    
    # Проверяем, что это Администрация
    if org.Логин != 'Administracia':
        messages.error(request, 'Этот раздел доступен только для администрации')
        return redirect('organization_dashboard')
    
    smeti = СметаЗатрат.objects.filter(организация=org).order_by('-дата_отправки')
    
    # Статистика
    stats = {
        'черновики': smeti.filter(статус='черновик').count(),
        'отправленные': smeti.filter(статус='отправлена').count(),
        'на_рассмотрении': smeti.filter(статус='на_рассмотрении').count(),
        'одобренные': smeti.filter(статус='одобрена').count(),
        'отклоненные': smeti.filter(статус='отклонена').count(),
        'на_доработке': smeti.filter(статус='требует_доработки').count(),
    }
    
    context = {
        'organization': org,
        'smeti': smeti,
        'stats': stats,
    }
    return render(request, 'index/organization_smeti_list.html', context)


def organization_smeta_create(request):
    """Создание новой сметы администрацией"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    org_id = request.session.get('organization_id')
    try:
        org = Организация.objects.get(id=org_id)
    except Организация.DoesNotExist:
        return redirect('organization_logout')
    
    # Проверяем, что это Администрация
    if org.Логин != 'Administracia':
        messages.error(request, 'Этот раздел доступен только для администрации')
        return redirect('organization_dashboard')
    
    if request.method == 'POST':
        form = СметаЗатратФорма(request.POST, request.FILES)
        if form.is_valid():
            smeta = form.save(commit=False)
            smeta.организация = org
            smeta.статус = 'черновик'
            smeta.save()
            
            messages.success(request, 'Смета успешно создана')
            # ИЗМЕНЕНО: перенаправляем на детальную страницу, а не на редактирование проектов
            return redirect('organization_smeta_detail', pk=smeta.pk)
    else:
        form = СметаЗатратФорма(initial={'год': timezone.now().year + 1})
    
    context = {
        'organization': org,
        'form': form,
    }
    return render(request, 'index/organization_smeta_create.html', context)

def organization_smeta_detail(request, pk):
    """Детальный просмотр сметы для администрации"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    org_id = request.session.get('organization_id')
    try:
        org = Организация.objects.get(id=org_id)
    except Организация.DoesNotExist:
        return redirect('organization_logout')
    
    smeta = get_object_or_404(СметаЗатрат, pk=pk, организация=org)
    
    # Форма для добавления проекта
    if request.method == 'POST' and 'add_project' in request.POST:
        project_form = СметаПроектФорма(request.POST)
        if project_form.is_valid():
            project = project_form.save()
            smeta.проекты.add(project)
            messages.success(request, f'Проект "{project.название}" добавлен в смету')
            return redirect('organization_smeta_detail', pk=smeta.pk)
    else:
        project_form = СметаПроектФорма()
    
    context = {
        'organization': org,
        'smeta': smeta,
        'project_form': project_form,
        'can_edit': smeta.статус == 'черновик',
        'can_send': smeta.статус == 'черновик' and smeta.проекты.exists(),
    }
    return render(request, 'index/organization_smeta_detail.html', context)

def organization_smeta_send(request, pk):
    """Отправка сметы в бюджетный отдел"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    org_id = request.session.get('organization_id')
    try:
        org = Организация.objects.get(id=org_id)
    except Организация.DoesNotExist:
        return redirect('organization_logout')
    
    smeta = get_object_or_404(СметаЗатрат, pk=pk, организация=org)
    
    if smeta.статус not in ['черновик', 'требует_доработки']:
        messages.error(request, 'Эту смету нельзя отправить')
        return redirect('organization_smeta_detail', pk=smeta.pk)
    
    if smeta.проекты.count() == 0:
        messages.error(request, 'Добавьте хотя бы один проект в смету перед отправкой')
        return redirect('organization_smeta_detail', pk=smeta.pk)
    
    smeta.статус = 'отправлена'
    smeta.дата_отправки = timezone.now()
    smeta.save()
    
    messages.success(request, 'Смета успешно отправлена в бюджетный отдел')
    return redirect('organization_smeta_detail', pk=smeta.pk)


def organization_smeta_edit_projects(request, pk):
    """Редактирование проектов в смете (для администрации)"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    org_id = request.session.get('organization_id')
    try:
        org = Организация.objects.get(id=org_id)
    except Организация.DoesNotExist:
        return redirect('organization_logout')
    
    smeta = get_object_or_404(СметаЗатрат, pk=pk, организация=org)
    
    if smeta.статус not in ['черновик', 'требует_доработки']:
        messages.error(request, 'Нельзя редактировать отправленную смету')
        return redirect('organization_smeta_detail', pk=smeta.pk)
    
    projects = smeta.проекты.all()
    
    if request.method == 'POST':
        if 'delete_project' in request.POST:
            project_id = request.POST.get('project_id')
            project = get_object_or_404(СметаПроект, pk=project_id)
            smeta.проекты.remove(project)
            project.delete()
            messages.success(request, 'Проект удален')
            return redirect('organization_smeta_edit_projects', pk=smeta.pk)
    
    context = {
        'organization': org,
        'smeta': smeta,
        'projects': projects,
    }
    return render(request, 'index/organization_smeta_edit_projects.html', context)


# ==================== БЮДЖЕТНЫЙ ОТДЕЛ: РАБОТА СО СМЕТАМИ ====================

@custom_user_required
def budget_smeti_list(request):
    """Отдельная страница для смет в бюджетном отделе - только отправленные"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    # Получаем только отправленные сметы (не черновики)
    smeti = СметаЗатрат.objects.filter(статус='отправлена').order_by('-дата_отправки')
    
    context = {
        'user': user,
        'smeti': smeti,
    }
    return render(request, 'index/budget_smeti_list.html', context)


@custom_user_required
def budget_smeta_detail(request, pk):
    """Детальный просмотр сметы для бюджетного отдела"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    smeta = get_object_or_404(СметаЗатрат, pk=pk)
    
    context = {
        'user': user,
        'smeta': smeta,
        'can_review': smeta.статус in ['отправлена', 'на_рассмотрении'],
        'can_create_applications': smeta.статус == 'одобрена',
    }
    return render(request, 'index/budget_smeta_detail.html', context)


@custom_user_required
def budget_smeta_review(request, pk):
    """Рассмотрение сметы бюджетным отделом"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    smeta = get_object_or_404(СметаЗатрат, pk=pk)
    
    # Проверяем статус - только отправленные сметы можно рассматривать
    if smeta.статус not in ['отправлена']:
        messages.error(request, 'Эту смету уже рассмотрели')
        return redirect('budget_smeti_list')  # ← ИЗМЕНИЛ на список
    
    if request.method == 'POST':
        # Получаем действие из POST данных (не из формы)
        действие = request.POST.get('действие')
        комментарий = request.POST.get('комментарий', '')
        
        if действие == 'approve':
            smeta.одобрить(user, комментарий)
            messages.success(request, f'✅ Смета "{smeta.название}" одобрена')
        elif действие == 'reject':
            smeta.отклонить(user, комментарий)
            messages.error(request, f'❌ Смета "{smeta.название}" отклонена')
        else:
            messages.warning(request, 'Неизвестное действие')
        
        # ВАЖНО: перенаправляем обратно на список смет
        return redirect('budget_smeti_list')
    
    # Для GET запроса показываем форму (если нужно)
    context = {
        'user': user,
        'smeta': smeta,
    }
    return render(request, 'index/budget_smeta_review.html', context)


@custom_user_required
def budget_create_applications_from_smeta(request, pk):
    """Создание заявок из одобренной сметы"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    smeta = get_object_or_404(СметаЗатрат, pk=pk, статус='одобрена')
    
    if request.method == 'POST':
        form = СоздатьЗаявкиИзСметыФорма(request.POST, смета=smeta)
        if form.is_valid():
            выбранные_проекты = form.cleaned_data['проекты']
            созданные_заявки = []
            
            for проект in выбранные_проекты:
                # Создаем заявку для каждого проекта
                заявка = Заявка.objects.create(
                    организация=smeta.организация,
                    полное_наименование=smeta.организация.ПолноеНаименование,
                    инн=smeta.организация.ИНН,
                    счет_организации=smeta.организация.СчетОрганизации,
                    вид_выплаты=None,  # Можно выбрать или создать специальный вид
                    другой_вид_выплаты=f"Смета: {проект.название}",
                    запрашиваемая_сумма=проект.запрашиваемая_сумма,
                    статус='new'
                )
                созданные_заявки.append(заявка)
            
            # Сохраняем связь
            smeta.созданные_заявки.add(*созданные_заявки)
            
            messages.success(
                request, 
                f'Создано {len(созданные_заявки)} заявок из сметы. Они появились в списке новых заявок.'
            )
            return redirect('budget_smeta_detail', pk=smeta.pk)
    else:
        form = СоздатьЗаявкиИзСметыФорма(смета=smeta)
    
    context = {
        'user': user,
        'smeta': smeta,
        'form': form,
    }
    return render(request, 'index/budget_create_applications_from_smeta.html', context)


def smeta_download_document(request, pk):
    """Скачивание документа сметы"""
    smeta = get_object_or_404(СметаЗатрат, pk=pk)
    
    # Проверяем права доступа
    is_organization = request.session.get('organization_authenticated')
    is_budget_user = request.session.get('user_id')
    
    if is_organization:
        org_id = request.session.get('organization_id')
        if org_id != smeta.организация_id:
            return HttpResponseForbidden('Нет доступа к этому документу')
    elif is_budget_user:
        try:
            user = Пользователь.objects.get(pk=is_budget_user)
            if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
                return HttpResponseForbidden('Нет доступа к этому документу')
        except:
            return HttpResponseForbidden('Нет доступа к этому документу')
    else:
        return redirect('welcome')
    
    if not smeta.документ:
        messages.error(request, 'Документ не найден')
        return redirect(request.META.get('HTTP_REFERER', 'welcome'))
    
    # Добавьте отладку
    print(f"Скачивание файла: {smeta.документ.path}")
    print(f"Файл существует: {os.path.exists(smeta.документ.path)}")
    
    try:
        with open(smeta.документ.path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/octet-stream')
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(smeta.документ.name)}"'
            return response
    except FileNotFoundError:
        messages.error(request, 'Файл не найден на диске')
        return redirect(request.META.get('HTTP_REFERER', 'welcome'))

def generate_official_letter(smeta, данные):
    """Генерация официального письма-заявки"""
    
    # Формируем текст документа
    документ = f"""
    Главе муниципального образования
    «Лениногорский муниципальный район»
    ___________________________
    
    от Главы администрации
    города Лениногорск
    ___________________________
    
    Уважаемый ___________________!
    
    В соответствии с Соглашением о сотрудничестве № {данные['номер_соглашения']} 
    от {данные['дата_соглашения'].strftime('%d.%m.%Y')} просим рассмотреть возможность 
    выделения средств {данные['наименование_учреждения']} 
    на реализацию следующих проектов в {smeta.год} году:
    
    """
    
    # Добавляем выбранные проекты
    общая_сумма = 0
    for проект in данные['проекты']:
        документ += f"""
    • {проект.название}
      Описание: {проект.описание or 'Не указано'}
      Запрашиваемая сумма: {проект.запрашиваемая_сумма:,.2f} руб.
      """
        общая_сумма += проект.запрашиваемая_сумма
    
    документ += f"""
    
    Общая сумма финансирования: {общая_сумма:,.2f} руб.
    
    {данные['дополнительный_текст']}
    
    Приложения:
    1. Смета расходов на {smeta.год} год (прилагается отдельным файлом)
    2. Экономическое обоснование по каждому проекту
    
    С уважением,
    
    Глава администрации города Лениногорск
    __________________ /__________________/
    
    М.П.
    
    Исполнитель: __________________
    Тел.: __________________
    """
    
    return документ


def смета_создать_документ(request, pk):
    """Создание документов для сметы"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    smeta = get_object_or_404(СметаЗатрат, pk=pk)
    
    if request.method == 'POST':
        form = СметаДанныеФорма(request.POST, smeta=smeta)
        if form.is_valid():
            data = form.cleaned_data
            
            # Обновляем общую сумму сметы перед генерацией
            total = 0
            for project in smeta.проекты.all():
                total += project.запрашиваемая_сумма
            smeta.общая_сумма = total
            smeta.save()
            
            # Генерируем PDF сметы (функция из utils.py)
            pdf_smeta = generate_smeta_pdf(smeta)
            
            # Сохраняем смету в БД
            doc_smeta = ДокументСметы.objects.create(
                смета=smeta,
                тип='smeta',
                название=f"Смета на {smeta.год} год",
                pdf_data=pdf_smeta.getvalue()
            )
            
            # Генерируем PDF заявки (функция из utils.py)
            selected_projects = data['проекты']
            pdf_zayavka = generate_zayavka_pdf(smeta, data, selected_projects)
            
            # Сохраняем заявку в БД
            doc_zayavka = ДокументСметы.objects.create(
                смета=smeta,
                тип='zayavka',
                название=f"Письмо-заявка от {datetime.now().strftime('%d.%m.%Y')}",
                pdf_data=pdf_zayavka.getvalue()
            )
            
            messages.success(request, 'Документы успешно созданы и прикреплены к смете')
            return redirect('organization_smeta_detail', pk=smeta.pk)
    else:
        form = СметаДанныеФорма(smeta=smeta)
    
    context = {
        'smeta': smeta,
        'form': form,
    }
    return render(request, 'index/smeta_create_document.html', context)

    
    response = HttpResponse(doc.pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{doc.название}.pdf"'
    return response


def download_smeta_pdf(request, pk):
    """Скачивание сметы в PDF с электронной печатью"""
    smeta = get_object_or_404(СметаЗатрат, pk=pk)
    
    # Проверка прав доступа
    if not (request.session.get('organization_authenticated') or 
            request.session.get('user_id')):
        return redirect('welcome')
    
    # Генерируем PDF
    pdf_buffer = generate_smeta_pdf(smeta)
    
    # Сохраняем документ в базу данных
    doc = ДокументСметы.objects.create(
        смета=smeta,
        тип='smeta',
        название=f"Смета на {smeta.год} год",
        pdf_data=pdf_buffer.getvalue()  # ИСПРАВЛЕНО: файл -> pdf_data
    )
    
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="smeta_{smeta.год}.pdf"'
    return response

def create_zayavka_pdf(request, pk):
    """Создание письма-заявки"""
    smeta = get_object_or_404(СметаЗатрат, pk=pk)
    
    if request.method == 'POST':
        # Собираем данные из формы
        form_data = {
            'глава': request.POST.get('glava', 'Иванову И.И.'),
            'должность_главы': request.POST.get('dolzhnost', 'Главе Лениногорского муниципального района'),
            'от_кого': request.POST.get('ot_kogo', 'Главы администрации города Лениногорск'),
            'фио_отправителя': request.POST.get('fio', 'Петрова П.П.'),
            'номер_соглашения': request.POST.get('dogovor_num', '45-С/2025'),
            'дата_соглашения': request.POST.get('dogovor_date', datetime.now().strftime('%d.%m.%Y')),
            'основание': request.POST.get('osnovanie', 'в соответствии с Соглашением о сотрудничестве'),
            'учреждение': request.POST.get('uchrezhdenie', 'Муниципальному казённому учреждению «Финансово-бюджетная палата»'),
            'проекты': smeta.проекты.all(),
            'итого': smeta.общая_сумма,
            'доп_текст': request.POST.get('dop_text', ''),
            'исполнитель': request.POST.get('ispolnitel', 'Сидорова А.А.'),
            'телефон': request.POST.get('telefon', '8 (85595) 5-00-00'),
        }
        
        # Генерируем PDF
        pdf_buffer = generate_zayavka_pdf(smeta, form_data)
        
        # Сохраняем документ
        doc = ДокументСметы.objects.create(
            смета=smeta,
            тип='zayavka',
            название=f"Письмо-заявка от {datetime.now().strftime('%d.%m.%Y')}",
            файл=None
        )
        
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="zayavka_{smeta.id}.pdf"'
        return response
    
    # Показываем форму для заполнения
    context = {
        'smeta': smeta,
        'now': datetime.now(),
    }
    return render(request, 'index/create_zayavka.html', context)

    
def заявка_загрузить_документ(request, pk):
    """Загрузка документа к заявке (для организации)"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    заявка = get_object_or_404(Заявка, pk=pk)
    org_id = request.session.get('organization_id')
    
    # Проверяем, что заявка принадлежит этой организации
    if заявка.организация_id != org_id:
        return HttpResponseForbidden('Нет доступа к этой заявке')
    
    # Проверяем, что заявка ещё новая (можно добавлять документы только к новым)
    if заявка.статус != 'new':
        messages.error(request, 'Нельзя добавлять документы к обработанной заявке')
        return redirect('organization_applications')
    
    if request.method == 'POST':
        form = ДокументЗаявкиФорма(request.POST, request.FILES)
        if form.is_valid():
            документ = form.save(commit=False)
            документ.заявка = заявка
            документ.save()
            messages.success(request, 'Документ успешно загружен')
            return redirect('organization_application_detail', pk=заявка.pk)
    else:
        form = ДокументЗаявкиФорма()
    
    context = {
        'заявка': заявка,
        'form': form,
    }
    return render(request, 'index/upload_document.html', context)

def заявка_удалить_документ(request, pk):
    """Удаление документа (только для организации-владельца)"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')
    
    документ = get_object_or_404(ДокументЗаявки, pk=pk)
    заявка = документ.заявка
    org_id = request.session.get('organization_id')
    
    # Проверяем права
    if заявка.организация_id != org_id:
        return HttpResponseForbidden('Нет доступа')
    
    if заявка.статус != 'new':
        messages.error(request, 'Нельзя удалять документы из обработанной заявки')
        return redirect('organization_application_detail', pk=заявка.pk)
    
    if request.method == 'POST':
        документ.delete()
        messages.success(request, 'Документ удалён')
    
    return redirect('organization_application_detail', pk=заявка.pk)

def скачать_документ(request, pk):
    """Скачивание документа сметы (включая письмо-заявку)"""
    try:
        doc = get_object_or_404(ДокументСметы, pk=pk)
        
        # Проверка прав доступа
        user_id = request.session.get('user_id')
        org_id = request.session.get('organization_id')
        
        print(f"=== СКАЧИВАНИЕ ДОКУМЕНТА СМЕТЫ ===")
        print(f"ID документа: {doc.id}, Тип: {doc.тип}, Название: {doc.название}")
        print(f"user_id: {user_id}, org_id: {org_id}")
        
        # Если это сотрудник бюджетного отдела
        if user_id:
            try:
                user = Пользователь.objects.get(pk=user_id)
                if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
                    print("❌ Нет доступа: не бюджетный отдел")
                    return HttpResponseForbidden('Нет доступа')
            except Пользователь.DoesNotExist:
                print("❌ Пользователь не найден")
                return HttpResponseForbidden('Нет доступа')
        # Если это организация (администрация)
        elif org_id:
            # Проверяем, что смета принадлежит этой организации
            if doc.смета.организация_id != org_id:
                print(f"❌ Нет доступа: организация {org_id} не владеет сметой {doc.смета_id}")
                return HttpResponseForbidden('Нет доступа')
            print(f"✅ Доступ разрешен для организации {org_id}")
        else:
            print("❌ Не авторизован - перенаправляем на вход организации")
            return redirect('organization_login')
        
        # Проверяем, что есть данные
        if not doc.pdf_data:
            print("❌ Нет данных PDF")
            return HttpResponse('Документ не содержит данных', status=404)
        
        # Отправляем файл
        response = HttpResponse(doc.pdf_data, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{doc.название}.pdf"'
        print(f"✅ Файл отправлен, размер: {len(doc.pdf_data)} байт")
        return response
        
    except Exception as e:
        print(f"Ошибка при скачивании документа: {e}")
        import traceback
        traceback.print_exc()
        messages.error(request, f'Ошибка при скачивании документа')
        return redirect(request.META.get('HTTP_REFERER', 'organization_smeti_list'))


def скачать_документ_заявки(request, pk):
    """Скачивание документа заявки"""
    print(f"=== СКАЧИВАНИЕ ДОКУМЕНТА ЗАЯВКИ ID: {pk} ===")
    
    документ = get_object_or_404(ДокументЗаявки, pk=pk)
    
    # Проверка прав доступа
    user_id = request.session.get('user_id')
    org_id = request.session.get('organization_id')
    
    print(f"user_id: {user_id}, org_id: {org_id}")
    
    # Если это сотрудник бюджетного отдела
    if user_id:
        try:
            user = Пользователь.objects.get(pk=user_id)
            if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
                print("❌ Нет доступа: не бюджетный отдел")
                return HttpResponseForbidden('Нет доступа')
        except Пользователь.DoesNotExist:
            print("❌ Пользователь не найден")
            return HttpResponseForbidden('Нет доступа')
    # Если это организация (администрация)
    elif org_id:
        print(f"Организация ID: {org_id}, заявка организации: {документ.заявка.организация_id}")
        if org_id != документ.заявка.организация_id:
            print("❌ Нет доступа: не та организация")
            return HttpResponseForbidden('Нет доступа')
    else:
        print("❌ Не авторизован")
        return redirect('organization_login')  # ← ПРАВИЛЬНО: на вход организаций!
    
    # Проверяем, есть ли файл
    if not документ.файл:
        print("❌ Нет файла в документе")
        return HttpResponse('Файл не найден', status=404)
    
    # Проверяем, существует ли файл на диске
    import os
    if not os.path.exists(документ.файл.path):
        print(f"❌ Файл не существует: {документ.файл.path}")
        return HttpResponse('Файл не найден на диске', status=404)
    
    # Открываем и отправляем файл
    with open(документ.файл.path, 'rb') as f:
        file_content = f.read()
    
    print(f"✅ Файл прочитан, размер: {len(file_content)} байт")
    
    response = HttpResponse(file_content, content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{os.path.basename(документ.файл.name)}"'
    return response

def organization_application_detail(request, pk):
    """Детальный просмотр заявки для организации"""
    if not request.session.get('organization_authenticated'):
        return redirect('organization_login')

    org_id = request.session.get('organization_id')
    заявка = get_object_or_404(Заявка, pk=pk, организация_id=org_id)

    context = {
        'заявка': заявка,
    }
    return render(request, 'index/organization_application_detail.html', context)

def бюджет_смета_на_основе_выплат(request):
    """Создание сметы ФБП на основе выплаченных заявок от муниципальных организаций"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    # Получаем все года, за которые есть ВЫПЛАЧЕННЫЕ заявки
    годы_с_выплатами = Выплаты.objects.filter(статус='выплачено').dates('дата_выполнения', 'year').distinct()
    годы = sorted([г.year for г in годы_с_выплатами], reverse=True)
    
    if request.method == 'POST':
        form = ВыборГодаДляСметыФорма(request.POST, годы=годы)
        if form.is_valid():
            выбранный_год = int(form.cleaned_data['год'])
            следующий_год = выбранный_год + 1
            
            # Создаем новую смету ФБП на СЛЕДУЮЩИЙ год
            смета = СметаФБП.objects.create(
                год=следующий_год,
                создатель=user,
                статус='черновик',
                создана_на_основе_выплат_за_год=выбранный_год
            )
            
            # Анализируем выплаты за выбранный год (только выплаченные)
            выплаты = Выплаты.objects.filter(
                дата_выполнения__year=выбранный_год,
                статус='выплачено'
            ).select_related('вид', 'заявка__организация')
            
            # Группируем по видам выплат и организациям
            статьи_данные = {}
            
            for выплата in выплаты:
                # Определяем название статьи
                if выплата.вид:
                    название = выплата.вид.Название
                else:
                    название = 'Прочие выплаты'
                
                # Получаем организацию
                организация = выплата.заявка.организация.ПолноеНаименование if выплата.заявка else 'Неизвестно'
                
                ключ = f"{название}_{организация}"
                
                if ключ not in статьи_данные:
                    статьи_данные[ключ] = {
                        'название': название,
                        'организация': организация,
                        'сумма': 0,
                        'количество': 0
                    }
                
                статьи_данные[ключ]['сумма'] += выплата.сумма
                статьи_данные[ключ]['количество'] += 1
            
            # Создаем статьи сметы
            for данные in статьи_данные.values():
                СтатьяСметыФБП.objects.create(
                    смета=смета,
                    название=данные['название'],
                    тип='other',
                    сумма=данные['сумма'],
                    описание=f"На основе выплат {выбранный_год} года от организации {данные['организация']} ({данные['количество']} выплат)",
                    основано_на_выплатах=True,
                    количество_выплат=данные['количество'],
                    исходная_организация=данные['организация']
                )
            
            смета.рассчитать_общую_сумму()
            
            messages.success(request, f'Смета ФБП на {следующий_год} год создана на основе выплат {выбранный_год} года')
            return redirect('budget_smeta_fbp_detail', pk=смета.pk)
    else:
        form = ВыборГодаДляСметыФорма(годы=годы)
    
    context = {
        'user': user,
        'form': form,
        'годы': годы,
    }
    return render(request, 'index/budget_smeta_fbp_osnova_vyplat.html', context)

def бюджет_сметы_fbp_list(request):
    """Список смет ФБП"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    сметы = СметаФБП.objects.all().order_by('-год', '-дата_создания')
    
    stats = {
        'черновики': сметы.filter(статус='черновик').count(),
        'утверждены': сметы.filter(статус='утверждена').count(),
        'отправлены': сметы.filter(статус='отправлена').count(),
        'всего': сметы.count(),
    }
    
    context = {
        'user': user,
        'сметы': сметы,
        'stats': stats,
    }
    return render(request, 'index/budget_smeti_fbp_list.html', context)


def бюджет_смета_fbp_детали(request, pk):
    """Детальный просмотр сметы ФБП"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    смета = get_object_or_404(СметаФБП, pk=pk)
    
    context = {
        'user': user,
        'смета': смета,
        'can_edit': смета.статус == 'черновик',
    }
    return render(request, 'index/budget_smeta_fbp_detail.html', context)

def бюджет_смета_fbp_создать_документы(request, pk):
    """Создание документов для сметы ФБП и отправка на почту"""
    print(f"=== СОЗДАНИЕ ДОКУМЕНТОВ ДЛЯ СМЕТЫ {pk} ===")
    
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    смета = get_object_or_404(СметаФБП, pk=pk)
    print(f"Смета найдена: {смета.год} год, статус: {смета.статус}")
    
    try:
        # Генерируем PDF
        print("Генерация PDF сметы...")
        pdf_smeta = generate_fbp_smeta_pdf(смета, user)
        print(f"PDF сметы сгенерирован, размер: {len(pdf_smeta.getvalue())}")
        
        # Сохраняем документ сметы
        doc_smeta = ДокументСметыФБП.objects.create(
            смета=смета,
            тип='smeta',
            название=f"Смета ФБП на {смета.год} год от {timezone.now().strftime('%d.%m.%Y')}",
            pdf_data=pdf_smeta.getvalue()
        )
        print(f"Документ сметы сохранен ID: {doc_smeta.id}")
        
        # Генерируем PDF письма
        print("Генерация письма...")
        pdf_pismo = generate_fbp_pismo_pdf(смета, user)
        print(f"PDF письма сгенерирован, размер: {len(pdf_pismo.getvalue())}")
        
        doc_pismo = ДокументСметыФБП.objects.create(
            смета=смета,
            тип='pismo',
            название=f"Письмо-заявка на {смета.год} год",
            pdf_data=pdf_pismo.getvalue()
        )
        print(f"Документ письма сохранен ID: {doc_pismo.id}")
        
        # Меняем статус сметы на "отправлена"
        смета.статус = 'отправлена'
        смета.save()
        print(f"Статус сметы изменен на 'отправлена'")
        
        # Отправляем на ПОЧТУ (вместо Telegram)
        print("Отправка на почту...")
        отправлено = отправить_смету_на_почту(смета, pdf_smeta.getvalue(), pdf_pismo.getvalue())
        print(f"Результат отправки: {отправлено}")
        
        if отправлено:
            messages.success(request, '✅ Документы созданы и отправлены на почту! Смета отправлена на рассмотрение.')
        else:
            messages.warning(request, '⚠️ Документы созданы, но не отправлены на почту. Смета сохранена как отправленная.')
            
    except Exception as e:
        print(f"ОШИБКА: {str(e)}")
        import traceback
        traceback.print_exc()
        messages.error(request, f'❌ Ошибка: {str(e)}')
    
    return redirect('budget_smeta_fbp_detail', pk=смета.pk)

##def отправить_смету_в_телеграм(смета, pdf_smeta_data, pdf_pismo_data):
    """Отправка документов в Telegram бот"""
    print("=== НАЧАЛО ОТПРАВКИ В TELEGRAM ===")
    
    token = getattr(settings, 'TELEGRAM_BOT_TOKEN', None)
    chat_id = getattr(settings, 'TELEGRAM_CHAT_ID', None)
    
    print(f"Token: {token[:10]}... (если есть)")
    print(f"Chat ID: {chat_id}")
    
    if not token or not chat_id:
        print("❌ ОШИБКА: Нет токена или chat_id")
        return False
    
    try:
        import time
        
        # 1. Сначала отправляем сообщение
        text = f"📄 <b>Новая смета ФБП на {смета.год} год</b>\n\n"
        text += f"💰 <b>Общая сумма:</b> {смета.общая_сумма:,.2f} руб.\n"
        text += f"📊 <b>Статей:</b> {смета.статьи.count()}\n"
        text += f"👤 <b>Создатель:</b> {смета.создатель.Фамилия} {смета.создатель.Имя}\n\n"
        text += f"✅ <b>Смета одобрена!</b>"
        
        url = f"https://api.telegram.org/bot{token}/sendMessage"
        data = {
            'chat_id': chat_id,
            'text': text,
            'parse_mode': 'HTML'
        }
        print("📤 Отправка сообщения...")
        response = requests.post(url, data=data)
        print(f"Ответ на сообщение: {response.status_code}")
        print(f"Текст ответа: {response.text}")
        
        if response.status_code != 200:
            print(f"❌ Ошибка отправки сообщения: {response.text}")
            return False
        
        # Небольшая пауза между отправками
        time.sleep(1)
        
        # 2. Отправляем PDF сметы
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        files = {'document': (f'smeta_{смета.год}.pdf', pdf_smeta_data, 'application/pdf')}
        data = {'chat_id': chat_id}
        print("📤 Отправка PDF сметы...")
        response = requests.post(url, data=data, files=files)
        print(f"Ответ на PDF сметы: {response.status_code}")
        
        if response.status_code != 200:
            print(f"❌ Ошибка отправки PDF сметы: {response.text}")
        
        time.sleep(1)
        
        # 3. Отправляем PDF письма
        files = {'document': (f'pismo_{смета.год}.pdf', pdf_pismo_data, 'application/pdf')}
        print("📤 Отправка PDF письма...")
        response = requests.post(url, data=data, files=files)
        print(f"Ответ на PDF письма: {response.status_code}")
        
        if response.status_code != 200:
            print(f"❌ Ошибка отправки PDF письма: {response.text}")
        
        print("✅ ОТПРАВКА ЗАВЕРШЕНА")
        return True
        
    except Exception as e:
        print(f"❌ ОШИБКА В TELEGRAM: {str(e)}")
        import traceback
        traceback.print_exc()
        return False
    
def бюджет_смета_fbp_из_администрации(request):
    """Добавление сметы от администрации в план ФБП"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    # Получаем одобренные сметы от администрации
    сметы_администрации = СметаЗатрат.objects.filter(статус='одобрена').order_by('-год')
    
    if request.method == 'POST':
        смета_ид = request.POST.get('смета')
        if смета_ид:
            смета_адм = get_object_or_404(СметаЗатрат, pk=смета_ид)
            
            # Создаем новую смету ФБП
            смета = СметаФБП.objects.create(
                год=смета_адм.год,
                создатель=user,
                статус='черновик'
            )
            
            # Копируем проекты как статьи
            for проект in смета_адм.проекты.all():
                СтатьяСметыФБП.objects.create(
                    смета=смета,
                    название=проект.название,
                    тип='other',
                    сумма=проект.запрашиваемая_сумма,
                    описание=проект.описание,
                    исходная_организация='Администрация'
                )
            
            смета.рассчитать_общую_сумму()
            messages.success(request, f'✅ Смета от администрации добавлена в план ФБП')
            return redirect('budget_smeta_fbp_detail', pk=смета.pk)
        else:
            messages.error(request, '❌ Выберите смету')
    
    context = {
        'user': user,
        'сметы': сметы_администрации,
    }
    return render(request, 'index/budget_smeta_fbp_from_admin.html', context)

def бюджет_смета_fbp_добавить_администрацию(request, pk):
    """Добавление сметы от администрации в существующую смету ФБП"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    смета = get_object_or_404(СметаФБП, pk=pk)
    
    # Получаем одобренные сметы от администрации (кроме уже добавленных)
    существующие_админ_статьи = смета.статьи.filter(исходная_организация='Администрация')
    существующие_сметы_идс = []
    for статья in существующие_админ_статьи:
        # Пытаемся найти исходную смету по описанию
        pass  # Пока просто покажем все
    
    сметы_администрации = СметаЗатрат.objects.filter(статус='одобрена').order_by('-год')
    
    if request.method == 'POST':
        смета_ид = request.POST.get('смета')
        if смета_ид:
            смета_адм = get_object_or_404(СметаЗатрат, pk=смета_ид)
            
            # Копируем проекты как статьи
            for проект in смета_адм.проекты.all():
                СтатьяСметыФБП.objects.create(
                    смета=смета,
                    название=проект.название,
                    тип='other',
                    сумма=проект.запрашиваемая_сумма,
                    описание=проект.описание,
                    исходная_организация='Администрация'
                )
            
            смета.рассчитать_общую_сумму()
            messages.success(request, f'✅ Смета от администрации добавлена в текущую смету')
            return redirect('budget_smeta_fbp_detail', pk=смета.pk)
    
    context = {
        'user': user,
        'смета': смета,
        'сметы_администрации': сметы_администрации,
    }
    return render(request, 'index/budget_smeta_fbp_add_admin.html', context)

def бюджет_смета_fbp_удалить_статью(request, pk):
    """Удаление статьи из сметы ФБП"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    статья = get_object_or_404(СтатьяСметыФБП, pk=pk)
    смета = статья.смета
    
    статья.delete()
    смета.рассчитать_общую_сумму()
    
    messages.success(request, '✅ Статья удалена')
    return redirect('budget_smeta_fbp_detail', pk=смета.pk)

def download_fbp_document(request, pk):
    """Скачивание документа сметы ФБП"""
    try:
        doc = get_object_or_404(ДокументСметыФБП, pk=pk)
        
        # Проверка прав доступа
        user_id = request.session.get('user_id')
        if not user_id:
            return redirect('login')
        
        response = HttpResponse(doc.pdf_data, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{doc.название}.pdf"'
        return response
        
    except Exception as e:
        messages.error(request, f'Ошибка при скачивании документа: {e}')
        return redirect(request.META.get('HTTP_REFERER', 'budget_smeti_fbp'))
    
def budget_smeta_fbp_delete(request, pk):
    """Удаление сметы ФБП"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
        if not user.Отдел or user.Отдел.Название != 'Бюджетный отдел':
            return redirect('user_dashboard')
    except Пользователь.DoesNotExist:
        return redirect('login')
    
    smeta = get_object_or_404(СметаФБП, pk=pk)
    
    if request.method == 'POST':
        smeta.delete()
        messages.success(request, f'Смета на {smeta.год} год успешно удалена')
        return redirect('budget_smeti_fbp')
    
    return redirect('budget_smeti_fbp')

def organization_smeta_delete(request, pk):
    """Удаление сметы администрацией"""
    user_id = request.session.get('organization_id')
    if not user_id:
        return redirect('organization_login')
    
    organization = Организация.objects.get(pk=user_id)
    smeta = get_object_or_404(СметаЗатрат, pk=pk, организация=organization)
    
    if request.method == 'POST':
        smeta.delete()
        messages.success(request, f'✅ Смета "{smeta.название}" успешно удалена')
        return redirect('organization_smeti_list')
    
    return redirect('organization_smeti_list')

@custom_user_required
def create_application_report(request):
    """Создание отчета по заявкам на основе текущих фильтров (без сохранения в БД)"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    # Получаем параметры фильтрации из GET
    status_filter = request.GET.get('status', '')
    payment_type_filter = request.GET.get('payment_type', '')
    date_from_filter = request.GET.get('date_from', '')
    date_to_filter = request.GET.get('date_to', '')
    
    # Формируем запрос для заявок
    applications = Заявка.objects.all().order_by('-дата_подачи')
    
    # Применяем фильтры
    if status_filter:
        applications = applications.filter(статус=status_filter)
    
    if payment_type_filter:
        try:
            applications = applications.filter(вид_выплаты_id=int(payment_type_filter))
        except ValueError:
            pass
    
    if date_from_filter:
        try:
            date_from = datetime.strptime(date_from_filter, '%Y-%m-%d').date()
            applications = applications.filter(дата_подачи__date__gte=date_from)
        except ValueError:
            date_from_filter = None
    
    if date_to_filter:
        try:
            date_to = datetime.strptime(date_to_filter, '%Y-%m-%d').date()
            applications = applications.filter(дата_подачи__date__lte=date_to)
        except ValueError:
            date_to_filter = None
    
    # Генерируем Word документ (передаем 6 аргументов)
    doc_buffer = generate_application_report_word(
        applications, 
        user, 
        date_from_filter, 
        date_to_filter, 
        status_filter, 
        payment_type_filter
    )
    
    # Формируем имя файла
    filename = f'Отчет_по_заявкам_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    
    # Отправляем файл на скачивание
    response = HttpResponse(doc_buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response


def get_status_name(status):
    """Возвращает русское название статуса"""
    status_names = {
        'new': 'Новая',
        'approved': 'Принята',
        'rejected': 'Отклонена'
    }
    return status_names.get(status, status)


def generate_application_report_word(applications, user, date_from, date_to, status_filter, payment_type_filter):
    """Генерация официального отчета по заявкам в Word (без смайликов)"""
    
    doc = Document()
    
    # Настройка шрифта по умолчанию
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    # Верхний колонтитул (шапка) - центрирование
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    header_run = header_para.add_run('МУНИЦИПАЛЬНОЕ КАЗЕННОЕ УЧРЕЖДЕНИЕ\n')
    header_run.bold = True
    header_run.font.size = Pt(14)
    
    header_run2 = header_para.add_run('«ФИНАНСОВО-БЮДЖЕТНАЯ ПАЛАТА»\n')
    header_run2.bold = True
    header_run2.font.size = Pt(14)
    
    header_para.add_run('муниципального образования «Лениногорский муниципальный район»\n')
    header_para.add_run('Республики Татарстан\n\n')
    
    # Название документа - центрирование
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('ОТЧЕТ\n')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.add_run('о поступивших заявках от организаций\n\n')
    
    # Регистрационные данные
    doc.add_paragraph(f'Дата формирования отчета: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Время формирования: {datetime.now().strftime("%H:%M")}')
    doc.add_paragraph(f'Сформировал: {user.Фамилия} {user.Имя} {user.Отчество or ""}')
    doc.add_paragraph(f'Должность: {user.Отдел.Название if user.Отдел else "Сотрудник"}\n')
    
    # Раздел 1. Параметры фильтрации
    doc.add_paragraph('1. ПАРАМЕТРЫ ФИЛЬТРАЦИИ ДАННЫХ')
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph('')
    
    status_names = {
        'new': 'Новая',
        'approved': 'Принята',
        'rejected': 'Отклонена'
    }
    
    if date_from and date_to:
        period_text = f'1.1. Период рассмотрения: с {datetime.strptime(date_from, "%Y-%m-%d").strftime("%d.%m.%Y")} по {datetime.strptime(date_to, "%Y-%m-%d").strftime("%d.%m.%Y")}'
        doc.add_paragraph(period_text)
    elif date_from:
        period_text = f'1.1. Период рассмотрения: с {datetime.strptime(date_from, "%Y-%m-%d").strftime("%d.%m.%Y")} по настоящее время'
        doc.add_paragraph(period_text)
    elif date_to:
        period_text = f'1.1. Период рассмотрения: за период до {datetime.strptime(date_to, "%Y-%m-%d").strftime("%d.%m.%Y")}'
        doc.add_paragraph(period_text)
    else:
        doc.add_paragraph('1.1. Период рассмотрения: за все время')
    
    if status_filter:
        doc.add_paragraph(f'1.2. Статус заявок: {status_names.get(status_filter, status_filter)}')
    else:
        doc.add_paragraph('1.2. Статус заявок: все статусы')
    
    if payment_type_filter:
        try:
            payment = ВидыВыплат.objects.get(id=int(payment_type_filter))
            doc.add_paragraph(f'1.3. Вид выплаты: {payment.Название}')
        except:
            doc.add_paragraph(f'1.3. Вид выплаты: выбран конкретный вид')
    else:
        doc.add_paragraph('1.3. Вид выплаты: все виды')
    
    doc.add_paragraph('')
    
    # Подсчет статистики
    total_amount = applications.aggregate(total=Sum('запрашиваемая_сумма'))['total'] or 0
    new_count = applications.filter(статус='new').count()
    approved_count = applications.filter(статус='approved').count()
    rejected_count = applications.filter(статус='rejected').count()
    
    # Раздел 2. Итоговая статистика
    doc.add_paragraph('2. ИТОГОВАЯ СТАТИСТИКА ПО ЗАЯВКАМ')
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph('')
    
    doc.add_paragraph(f'2.1. Общее количество заявок, соответствующих критериям фильтрации: {applications.count()} шт.')
    doc.add_paragraph(f'2.2. Общая запрашиваемая сумма: {total_amount:,.2f} рублей')
    doc.add_paragraph(f'2.3. Количество заявок со статусом "Новая": {new_count} шт.')
    doc.add_paragraph(f'2.4. Количество заявок со статусом "Принята": {approved_count} шт.')
    doc.add_paragraph(f'2.5. Количество заявок со статусом "Отклонена": {rejected_count} шт.\n')
    
    # Раздел 3. Перечень заявок
    if applications.exists():
        doc.add_paragraph('3. ПЕРЕЧЕНЬ ЗАЯВОК')
        doc.paragraphs[-1].runs[0].bold = True
        doc.add_paragraph('')
        
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки таблицы
        headers = ['№ п/п', 'Регистрационный номер', 'Полное наименование организации', 'Вид выплаты', 'Запрашиваемая сумма (руб.)', 'Дата подачи заявки', 'Статус рассмотрения']
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
        
        # Заполняем таблицу данными
        for idx, app in enumerate(applications, 1):
            row = table.add_row()
            
            # Номер по порядку
            cell = row.cells[0]
            cell.text = str(idx)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # ID заявки
            cell = row.cells[1]
            cell.text = f'№ {app.id}'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Наименование организации
            org_name = app.полное_наименование[:60] if len(app.полное_наименование) > 60 else app.полное_наименование
            row.cells[2].text = org_name
            
            # Вид выплаты
            payment_name = app.вид_выплаты.Название if app.вид_выплаты else (app.другой_вид_выплаты or '—')
            row.cells[3].text = payment_name[:50] if len(payment_name) > 50 else payment_name
            
            # Сумма
            cell = row.cells[4]
            cell.text = f'{app.запрашиваемая_сумма:,.2f}'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Дата подачи
            cell = row.cells[5]
            cell.text = app.дата_подачи.strftime('%d.%m.%Y')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Статус
            status_text = {
                'new': 'Новая',
                'approved': 'Принята',
                'rejected': 'Отклонена'
            }.get(app.статус, app.статус)
            cell = row.cells[6]
            cell.text = status_text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем итоговую строку
        row = table.add_row()
        row.cells[0].text = ''
        row.cells[1].text = ''
        row.cells[2].text = ''
        row.cells[3].text = 'ИТОГО:'
        row.cells[4].text = f'{total_amount:,.2f}'
        row.cells[5].text = ''
        row.cells[6].text = ''
        
        # Выделяем итоговую строку жирным
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
        
        doc.add_paragraph('')
        doc.add_paragraph(f'Общая сумма заявок: {total_amount:,.2f} рублей.')
        doc.add_paragraph('')
        
    else:
        doc.add_paragraph('3. ПЕРЕЧЕНЬ ЗАЯВОК')
        doc.paragraphs[-1].runs[0].bold = True
        doc.add_paragraph('')
        doc.add_paragraph('Заявки, соответствующие критериям фильтрации, не обнаружены.')
        doc.paragraphs[-1].runs[0].bold = True
        doc.add_paragraph('')
    
    # Раздел 4. Заключение
    doc.add_paragraph('4. ЗАКЛЮЧЕНИЕ')
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph('')
    
    conclusion_para = doc.add_paragraph()
    if new_count > 0:
        conclusion_para.add_run(f'По состоянию на отчетную дату {new_count} заявок находятся на стадии рассмотрения. ')
    if approved_count > 0:
        conclusion_para.add_run(f'{approved_count} заявок одобрены и направлены для осуществления выплат. ')
    if rejected_count > 0:
        conclusion_para.add_run(f'{rejected_count} заявок отклонены по причинам, указанным в комментариях к заявкам. ')
    
    if applications.count() == 0:
        conclusion_para.add_run('За указанный период заявки от организаций не поступали.')
    
    doc.add_paragraph('')
    doc.add_paragraph('На основании вышеизложенного, предлагается:')
    doc.add_paragraph('1. Принять к сведению представленную информацию.')
    doc.add_paragraph('2. Продолжить мониторинг поступления заявок от организаций.')
    doc.add_paragraph('3. Обеспечить своевременное рассмотрение вновь поступающих заявок.\n')
    
    # Раздел 5. Подписи
    doc.add_paragraph('5. ПОДПИСИ')
    doc.paragraphs[-1].runs[0].bold = True
    doc.add_paragraph('')
    
    doc.add_paragraph('Ответственный сотрудник:')
    doc.add_paragraph(f'{user.Фамилия} {user.Имя} {user.Отчество or ""}')
    doc.add_paragraph('__________________________')
    doc.add_paragraph('(подпись)')
    
    doc.add_paragraph('')
    doc.add_paragraph('Согласовано:')
    doc.add_paragraph('Начальник отдела')
    doc.add_paragraph('__________________________')
    doc.add_paragraph('(подпись)')
    
    doc.add_paragraph('')
    doc.add_paragraph('Место печати')
    doc.add_paragraph('М.П.')
    
    # Сохраняем в буфер
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


@custom_user_required
def application_reports_list(request):
    """Список отчетов по заявкам"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    reports = ОтчетПоЗаявкам.objects.filter(создатель=user).order_by('-дата_создания')
    
    # Получаем параметры для отображения
    for report in reports:
        if report.фильтр_статус:
            status_names = {'new': 'Новая', 'approved': 'Принята', 'rejected': 'Отклонена'}
            report.status_display = status_names.get(report.фильтр_статус, report.фильтр_статус)
        else:
            report.status_display = 'Все статусы'
        
        if report.фильтр_вид_выплаты:
            try:
                payment = ВидыВыплат.objects.get(id=report.фильтр_вид_выплаты)
                report.payment_display = payment.Название
            except:
                report.payment_display = 'Выбран конкретный вид'
        else:
            report.payment_display = 'Все виды'
    
    context = {
        'user': user,
        'reports': reports,
    }
    return render(request, 'index/application_reports_list.html', context)


@custom_user_required
def download_application_report(request, pk):
    """Скачивание отчета по заявкам"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    report = get_object_or_404(ОтчетПоЗаявкам, pk=pk, создатель=user)
    
    if report.файл and report.файл.path:
        # Проверяем существует ли файл
        import os
        if os.path.exists(report.файл.path):
            with open(report.файл.path, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{report.название}.docx"'
                return response
        else:
            messages.error(request, 'Файл отчета не найден на сервере')
            return redirect('application_reports_list')
    else:
        messages.error(request, 'Файл отчета отсутствует')
        return redirect('application_reports_list')


@custom_user_required
def delete_application_report(request, pk):
    """Удаление отчета по заявкам"""
    user_id = request.session.get('user_id')
    try:
        user = Пользователь.objects.get(pk=user_id)
    except Пользователь.DoesNotExist:
        request.session.flush()
        return redirect('login')
    
    report = get_object_or_404(ОтчетПоЗаявкам, pk=pk, создатель=user)
    
    if request.method == 'POST':
        # Удаляем файл с сервера
        if report.файл and report.файл.path:
            import os
            if os.path.exists(report.файл.path):
                os.remove(report.файл.path)
        report.delete()
        messages.success(request, 'Отчет успешно удален')
        return redirect('application_reports_list')
    
    return redirect('application_reports_list')